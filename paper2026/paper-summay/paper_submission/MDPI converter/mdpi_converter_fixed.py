#!/usr/bin/env python3
"""
Fixed MDPI Markdown to Word Converter
--------------------------------------
Properly handles:
- Title extraction from YAML front matter
- Figure insertion from figure directory
- MDPI-compliant formatting
"""

import os
import sys
import re
import shutil
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE


def extract_yaml_frontmatter(content):
    """Extract YAML front matter from markdown content."""
    if not content.startswith('---'):
        return None, content
    
    # Find end of YAML
    lines = content.split('\n')
    yaml_end = -1
    for i in range(1, len(lines)):
        if lines[i].strip() == '---':
            yaml_end = i
            break
    
    if yaml_end == -1:
        return None, content
    
    # Extract YAML content
    yaml_content = '\n'.join(lines[1:yaml_end])
    remaining_content = '\n'.join(lines[yaml_end+1:])
    
    try:
        yaml_data = yaml.safe_load(yaml_content)
        return yaml_data, remaining_content
    except:
        return None, content


def setup_mdpi_styles(doc):
    """Set up MDPI formatting."""
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(1.75)
        section.right_margin = Cm(1.75)
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
    
    # Configure Normal style
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def insert_figure(doc, fig_num, figures_dir, output_figures_dir):
    """Insert a figure into the document."""
    figure_mapping = {
        1: '01_sample_distribution.png',
        2: '02_missing_data_heatmap.png',
        3: '06_personality_dimensions.png',
        4: '07_personality_heatmap.png',
        5: '03_performance_comparison.png',
        6: '04_effect_sizes.png',
        7: '05_percentage_improvement.png',
        8: '08_weighted_scores.png',
        9: '09_total_score_boxplot.png'
    }
    
    filename = figure_mapping.get(fig_num)
    if not filename:
        return False
    
    figure_path = Path(figures_dir) / filename
    if not figure_path.exists():
        print(f"⚠️  Warning: Figure file not found: {figure_path}")
        return False
    
    # Copy figure to output directory
    output_fig = Path(output_figures_dir) / filename
    shutil.copy2(figure_path, output_fig)
    
    # Insert figure
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(figure_path), width=Inches(6.0))
    
    print(f"✓ Inserted Figure {fig_num}: {filename}")
    return True


def add_formatted_text(paragraph, text):
    """Add text with inline markdown formatting."""
    # Simple pattern matching for bold, italic, code
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def process_abstract(abstract_text):
    """Process abstract as single paragraph."""
    # Remove bold markers from section labels
    abstract_text = re.sub(r'\*\*([Bb]ackground|[Oo]bjective|[Mm]ethods|[Rr]esults|[Cc]onclusions?)\*\*:', r'\1:', abstract_text)
    # Join lines
    abstract_text = ' '.join(abstract_text.split('\n'))
    return abstract_text.strip()


def add_table(doc, table_lines):
    """Add MDPI-formatted table."""
    if not table_lines or len(table_lines) < 2:
        return
    
    # Parse table
    rows = []
    for line in table_lines:
        # Skip separator lines
        if re.match(r'^\s*\|[\s\-:|]+\|\s*$', line):
            continue
        # Extract cells
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            # Clean markdown
            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
            cell_text = re.sub(r'`(.+?)`', r'\1', cell_text)
            cell.text = cell_text
            
            # Format
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                    if i == 0:  # Header row
                        run.bold = True
    
    doc.add_paragraph()  # Spacing after table


def convert_markdown_to_docx(md_file, output_file, figures_dir=None):
    """Main conversion function."""
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Extract YAML and get title
    yaml_data, content = extract_yaml_frontmatter(content)
    title = yaml_data.get('title', 'Untitled Document') if yaml_data else 'Untitled Document'
    
    # Create document
    doc = Document()
    setup_mdpi_styles(doc)
    
    # Add title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(16)
        run.bold = True
    
    print(f"✓ Added title: {title[:60]}...")
    
    # Setup output directories
    output_dir = Path(output_file).parent
    figures_output = output_dir / "figures"
    figures_output.mkdir(parents=True, exist_ok=True)
    
    # Process content
    lines = content.split('\n')
    i = 0
    in_abstract = False
    abstract_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle Abstract
        if re.match(r'^##\s+Abstract', line, re.IGNORECASE):
            in_abstract = True
            doc.add_heading('Abstract', level=2)
            abstract_lines = []
            i += 1
            continue
        
        if in_abstract:
            if line.strip().startswith('#'):
                # End of abstract
                abstract_text = '\n'.join(abstract_lines)
                abstract_text = process_abstract(abstract_text)
                para = doc.add_paragraph(abstract_text)
                para.style = 'Normal'
                in_abstract = False
                continue  # Process current line as heading
            elif line.strip():
                abstract_lines.append(line.strip())
            i += 1
            continue
        
        # Handle figure placeholders
        if re.match(r'^\*\*\[Figure\s+\d+\s+near\s+here\]\*\*', line, re.IGNORECASE):
            # Extract figure number
            match = re.search(r'Figure\s+(\d+)', line)
            if match and figures_dir:
                fig_num = int(match.group(1))
                
                # Look ahead for caption
                caption_text = ""
                if i + 2 < len(lines) and lines[i+2].startswith('**Figure'):
                    caption_text = lines[i+2]
                    # Add caption
                    caption_para = doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Extract clean caption text
                    clean_caption = re.sub(r'^\*\*Figure\s+\d+\.\*\*\s*', '', caption_text)
                    
                    # Add bold figure label
                    bold_run = caption_para.add_run(f"Figure {fig_num}. ")
                    bold_run.bold = True
                    bold_run.font.size = Pt(10)
                    bold_run.font.name = 'Times New Roman'
                    
                    # Add caption text
                    caption_run = caption_para.add_run(clean_caption)
                    caption_run.font.size = Pt(10)
                    caption_run.font.name = 'Times New Roman'
                    
                    # Insert the actual figure BEFORE the caption
                    # Get the caption paragraph index
                    para_index = len(doc.paragraphs) - 1
                    # Insert figure before caption
                    if insert_figure(doc, fig_num, figures_dir, figures_output):
                        # Move caption after figure (it's already added)
                        pass
                    
                    i += 3  # Skip placeholder, blank line, and caption
                    continue
            
            i += 1
            continue
        
        # Handle headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\{.*?\}', '', text).strip()
            
            if level <= 3:
                doc.add_heading(text, level=level)
            else:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
            
            i += 1
            continue
        
        # Handle tables
        if line.strip().startswith('**Table'):
            # Extract table caption
            table_caption = re.sub(r'^\*\*Table\s+\d+\.\*\*\s*', '', line.strip())
            para = doc.add_paragraph()
            bold_run = para.add_run(table_caption.strip('*'))
            bold_run.bold = True
            i += 1
            continue
        
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            add_table(doc, table_lines)
            table_lines = []
            continue
        
        # Handle regular paragraphs
        if line.strip() and not re.match(r'^[\-\*\_]{3,}$', line.strip()):
            para = doc.add_paragraph()
            add_formatted_text(para, line.strip())
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"\n✓ Document saved: {output_file}")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Convert Markdown to MDPI Word format')
    parser.add_argument('input_md', help='Input Markdown file')
    parser.add_argument('-o', '--output', default='output.docx', help='Output Word file')
    parser.add_argument('-f', '--figures', help='Directory containing figures')
    
    args = parser.parse_args()
    
    print("=" * 70)
    print("MDPI Markdown to Word Converter (Fixed)")
    print("=" * 70)
    print()
    
    convert_markdown_to_docx(args.input_md, args.output, args.figures)
    
    print()
    print("=" * 70)
    print("✓ Conversion Complete!")
    print("=" * 70)



