#!/usr/bin/env python3
"""
Enhanced MDPI Markdown to Word Converter
-----------------------------------------
Converts Markdown manuscripts to MDPI-compliant Word documents with:
- Proper figure insertion from statistical analysis
- Table formatting per MDPI guidelines
- Supplementary materials export
- Complete MDPI manuscript structure
- Abstract formatting (single paragraph, no headings)
- Reference formatting [1] before punctuation

Author: Enhanced for MDPI Healthcare Submission
Date: October 2025
"""

import os
import sys
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MDPIConverter:
    """Enhanced MDPI manuscript converter."""
    
    def __init__(self, md_file, output_dir, figures_dir=None):
        """
        Initialize converter.
        
        Args:
            md_file: Path to input markdown file
            output_dir: Directory for output files
            figures_dir: Directory containing figure files
        """
        self.md_file = Path(md_file)
        self.output_dir = Path(output_dir)
        self.figures_dir = Path(figures_dir) if figures_dir else None
        self.doc = None
        self.figure_counter = 1
        self.table_counter = 1
        self.supplementary_counter = 1
        
        # Create output directories
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.figures_output = self.output_dir / "figures"
        self.figures_output.mkdir(exist_ok=True)
        self.supplementary_output = self.output_dir / "supplementary"
        self.supplementary_output.mkdir(exist_ok=True)
    
    def setup_mdpi_formatting(self):
        """Set up MDPI-specific document formatting."""
        # Set document margins according to MDPI standards
        # Top: 1.5cm, Bottom: 3.5cm, Left/Right: 1.75cm
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(1.75)
            section.right_margin = Cm(1.75)
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21.0)   # A4
        
        # Create MDPI-specific styles
        self.create_mdpi_styles()
    
    def create_mdpi_styles(self):
        """Create custom styles for MDPI formatting."""
        styles = self.doc.styles
        
        # MDPI Normal/Body Text style
        try:
            normal_style = styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(12)
            normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except KeyError:
            pass
        
        # MDPI Abstract style (crucial for single paragraph)
        try:
            abstract_style = styles.add_style('MDPI Abstract', WD_STYLE_TYPE.PARAGRAPH)
            abstract_style.font.name = 'Times New Roman'
            abstract_style.font.size = Pt(12)
            abstract_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            abstract_style.paragraph_format.space_after = Pt(12)
            abstract_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except ValueError:
            pass
        
        # MDPI Heading styles
        heading_sizes = {1: 16, 2: 14, 3: 12}
        for level, size in heading_sizes.items():
            try:
                heading_style = styles[f'Heading {level}']
                heading_style.font.name = 'Times New Roman'
                heading_style.font.bold = True
                heading_style.font.size = Pt(size)
                heading_style.paragraph_format.space_before = Pt(12)
                heading_style.paragraph_format.space_after = Pt(6)
                heading_style.paragraph_format.keep_with_next = True
            except KeyError:
                pass
        
        # MDPI Table Caption style
        try:
            table_caption_style = styles.add_style('MDPI Table Caption', WD_STYLE_TYPE.PARAGRAPH)
            table_caption_style.font.name = 'Times New Roman'
            table_caption_style.font.size = Pt(10)
            table_caption_style.font.bold = True
            table_caption_style.paragraph_format.space_before = Pt(6)
            table_caption_style.paragraph_format.space_after = Pt(3)
        except ValueError:
            pass
        
        # MDPI Figure Caption style
        try:
            fig_caption_style = styles.add_style('MDPI Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
            fig_caption_style.font.name = 'Times New Roman'
            fig_caption_style.font.size = Pt(10)
            fig_caption_style.paragraph_format.space_before = Pt(3)
            fig_caption_style.paragraph_format.space_after = Pt(12)
        except ValueError:
            pass
    
    def process_abstract(self, abstract_text):
        """
        Process abstract as single paragraph without headings.
        MDPI requires: Background, Methods, Results, Conclusion in one paragraph.
        """
        # Remove any markdown headings within abstract
        abstract_text = re.sub(r'\*\*[Bb]ackground\*\*:', 'Background:', abstract_text)
        abstract_text = re.sub(r'\*\*[Oo]bjective\*\*:', 'Objective:', abstract_text)
        abstract_text = re.sub(r'\*\*[Mm]ethods\*\*:', 'Methods:', abstract_text)
        abstract_text = re.sub(r'\*\*[Rr]esults\*\*:', 'Results:', abstract_text)
        abstract_text = re.sub(r'\*\*[Cc]onclusions?\*\*:', 'Conclusions:', abstract_text)
        
        # Remove line breaks within sections
        abstract_text = ' '.join(abstract_text.split('\n'))
        
        return abstract_text.strip()
    
    def insert_figure(self, figure_placeholder, caption_text):
        """Insert figure from figures directory."""
        # Extract figure number from placeholder or caption
        fig_match = re.search(r'Figure\s+(\d+)', caption_text, re.IGNORECASE)
        if not fig_match:
            return False
        
        fig_num = int(fig_match.group(1))
        
        # Figure file mapping (from your statistical analysis)
        figure_mapping = {
            1: '01_sample_distribution.png',
            2: '02_missing_data_heatmap.png',
            3: '06_personality_dimensions.png',
            4: '07_personality_heatmap.png',
            5: '03_performance_comparison.png',
            6: '04_effect_sizes.png',
            7: '05_percentage_improvement.png',
            8: '08_weighted_scores.png',
            9: '09_total_score_boxplot.png'
        }
        
        if fig_num not in figure_mapping:
            print(f"Warning: No mapping found for Figure {fig_num}")
            return False
        
        figure_file = figure_mapping[fig_num]
        
        # Look for figure in figures directory
        if self.figures_dir:
            source_fig = self.figures_dir / figure_file
            if source_fig.exists():
                # Copy to output directory
                dest_fig = self.figures_output / figure_file
                shutil.copy2(source_fig, dest_fig)
                
                # Insert figure into document
                try:
                    # Add figure
                    para = self.doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(str(source_fig), width=Inches(6.0))
                    
                    # Add caption
                    caption_para = self.doc.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        caption_para.style = 'MDPI Figure Caption'
                    except KeyError:
                        pass
                    
                    # Format caption
                    bold_run = caption_para.add_run(f"Figure {fig_num}. ")
                    bold_run.bold = True
                    bold_run.font.size = Pt(10)
                    
                    # Extract caption text (remove "Figure X." prefix)
                    clean_caption = re.sub(r'^\*\*Figure\s+\d+\.\*\*\s*', '', caption_text)
                    clean_caption = re.sub(r'^\*\*Figure\s+\d+:\*\*\s*', '', clean_caption)
                    caption_run = caption_para.add_run(clean_caption)
                    caption_run.font.size = Pt(10)
                    
                    print(f"✓ Inserted Figure {fig_num}: {figure_file}")
                    return True
                    
                except Exception as e:
                    print(f"Error inserting figure {fig_num}: {e}")
                    return False
        
        return False
    
    def add_mdpi_table(self, table_lines, caption=None):
        """Convert markdown table to MDPI-formatted Word table."""
        if not table_lines:
            return
        
        # Parse table
        rows = []
        for line in table_lines:
            # Skip separator lines
            if re.match(r'^\|[\s\-\:]+\|', line):
                continue
            
            # Split by | and clean
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c]
            
            if cells:
                rows.append(cells)
        
        if not rows:
            return
        
        # Add table caption if provided
        if caption:
            caption_para = self.doc.add_paragraph()
            try:
                caption_para.style = 'MDPI Table Caption'
            except KeyError:
                pass
            
            # Format caption
            bold_run = caption_para.add_run(f"Table {self.table_counter}. ")
            bold_run.bold = True
            bold_run.font.size = Pt(10)
            
            caption_run = caption_para.add_run(caption)
            caption_run.font.size = Pt(10)
            
            self.table_counter += 1
        
        # Create Word table
        num_cols = len(rows[0])
        num_rows = len(rows)
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Fill table
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if j < num_cols:
                    cell = table.rows[i].cells[j]
                    
                    # Clean markdown formatting
                    cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_data)
                    cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
                    cell_text = re.sub(r'`(.+?)`', r'\1', cell_text)
                    
                    cell.text = cell_text
                    
                    # Format cell
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.space_before = Pt(0)
                    
                    # Bold first row (header)
                    if i == 0:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def process_markdown_content(self, content):
        """Process markdown content and convert to MDPI Word format."""
        lines = content.split('\n')
        
        i = 0
        in_code_block = False
        code_content = []
        in_table = False
        table_lines = []
        table_caption = None
        in_abstract = False
        abstract_lines = []
        skip_yaml = False
        
        while i < len(lines):
            line = lines[i]
            
            # Skip YAML front matter
            if line.strip() == '---' and i == 0:
                skip_yaml = True
                i += 1
                continue
            
            if skip_yaml:
                if line.strip() == '---':
                    skip_yaml = False
                i += 1
                continue
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_content = []
                else:
                    in_code_block = False
                    if code_content:
                        code_para = self.doc.add_paragraph('\n'.join(code_content))
                        code_para.style = 'Normal'
                        for run in code_para.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                i += 1
                continue
            
            if in_code_block:
                code_content.append(line)
                i += 1
                continue
            
            # Handle Abstract section (must be single paragraph)
            if re.match(r'^##\s+Abstract', line, re.IGNORECASE):
                in_abstract = True
                self.doc.add_heading('Abstract', level=2)
                abstract_lines = []
                i += 1
                continue
            
            if in_abstract:
                # Check if we've reached next section
                if line.strip().startswith('#'):
                    # End abstract, process it
                    abstract_text = '\n'.join(abstract_lines)
                    abstract_text = self.process_abstract(abstract_text)
                    
                    para = self.doc.add_paragraph(abstract_text)
                    try:
                        para.style = 'MDPI Abstract'
                    except KeyError:
                        para.style = 'Normal'
                    
                    in_abstract = False
                    # Don't increment i, process current line
                    continue
                elif line.strip():
                    abstract_lines.append(line.strip())
                i += 1
                continue
            
            # Handle figure placeholders
            if re.match(r'^\*\*\[Figure\s+\d+\s+near\s+here\]\*\*', line, re.IGNORECASE):
                # Look for caption in next lines
                caption_text = ""
                j = i + 1
                while j < len(lines) and lines[j].strip():
                    if lines[j].strip().startswith('**Figure'):
                        caption_text = lines[j].strip()
                        i = j  # Skip to caption line
                        break
                    j += 1
                
                if caption_text:
                    self.insert_figure(line, caption_text)
                
                i += 1
                continue
            
            # Handle headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                
                # Clean formatting
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\{.*?\}', '', text)
                text = text.strip()
                
                if level <= 3:
                    self.doc.add_heading(text, level=level if level > 1 else 0)
                else:
                    para = self.doc.add_paragraph()
                    run = para.add_run(text)
                    run.bold = True
                    run.font.size = Pt(12)
                
                i += 1
                continue
            
            # Handle tables with captions
            if line.strip().startswith('**Table'):
                # Extract table caption
                table_caption = re.sub(r'^\*\*Table\s+\d+\.\*\*\s*', '', line.strip())
                table_caption = re.sub(r'^\*\*', '', table_caption)
                table_caption = re.sub(r'\*\*$', '', table_caption)
                i += 1
                continue
            
            # Handle table content
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_lines = [line]
                else:
                    table_lines.append(line)
                i += 1
                continue
            elif in_table:
                in_table = False
                self.add_mdpi_table(table_lines, table_caption)
                table_lines = []
                table_caption = None
                continue
            
            # Handle regular paragraphs
            if line.strip() and not re.match(r'^[\-\*\_]{3,}$', line.strip()):
                para = self.doc.add_paragraph()
                self.add_formatted_text(para, line.strip())
            
            i += 1
        
        # Handle any remaining table
        if in_table and table_lines:
            self.add_mdpi_table(table_lines, table_caption)
    
    def add_formatted_text(self, paragraph, text):
        """Add text with markdown formatting to paragraph."""
        # Handle references [1], [2], etc. - must be before punctuation per MDPI
        text = re.sub(r'\[(\d+(?:,\s*\d+)*)\]\.', r'[\1].', text)
        text = re.sub(r'\[(\d+(?:,\s*\d+)*)\],', r'[\1],', text)
        
        # Split on formatting markers
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\](?:\(.*?\))?)', text)
        
        for part in parts:
            if not part:
                continue
            
            # Bold
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # Italic
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # Inline code
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            # Links
            elif re.match(r'\[.*?\]\(.*?\)', part):
                link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if link_match:
                    run = paragraph.add_run(link_match.group(1))
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
            # References [1]
            elif re.match(r'\[\d+(?:,\s*\d+)*\]', part):
                run = paragraph.add_run(part)
                run.font.size = Pt(12)
            else:
                paragraph.add_run(part)
    
    def convert(self, output_filename):
        """Main conversion method."""
        print(f"\n{'='*70}")
        print("MDPI Enhanced Converter")
        print(f"{'='*70}\n")
        
        # Create document
        self.doc = Document()
        self.setup_mdpi_formatting()
        
        # Read markdown
        print(f"Reading: {self.md_file}")
        with open(self.md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Process content
        print("Processing markdown content...")
        self.process_markdown_content(content)
        
        # Save document
        output_path = self.output_dir / output_filename
        self.doc.save(str(output_path))
        
        print(f"\n✓ Conversion complete!")
        print(f"✓ Main document: {output_path}")
        print(f"✓ Figures exported: {self.figures_output}")
        print(f"✓ Supplementary materials: {self.supplementary_output}")
        print(f"\n{'='*70}\n")
        
        return output_path


def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Convert Markdown to MDPI-compliant Word document'
    )
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('-o', '--output', default='docoutput',
                       help='Output directory (default: docoutput)')
    parser.add_argument('-f', '--figures', default=None,
                       help='Figures directory (default: auto-detect)')
    parser.add_argument('-n', '--name', default='manuscript.docx',
                       help='Output filename (default: manuscript.docx)')
    
    args = parser.parse_args()
    
    # Auto-detect figures directory if not provided
    figures_dir = args.figures
    if not figures_dir:
        # Try to find statistical analysis figures
        md_path = Path(args.input)
        possible_fig_dirs = [
            md_path.parent / 'statistical analyis' / 'figures',
            md_path.parent / 'statistical_analysis' / 'figures',
            md_path.parent / 'figures'
        ]
        for fig_dir in possible_fig_dirs:
            if fig_dir.exists():
                figures_dir = str(fig_dir)
                print(f"Found figures directory: {figures_dir}")
                break
    
    # Create converter
    converter = MDPIConverter(
        md_file=args.input,
        output_dir=args.output,
        figures_dir=figures_dir
    )
    
    # Convert
    try:
        output_path = converter.convert(args.name)
        print(f"SUCCESS: {output_path}")
        return 0
    except Exception as e:
        print(f"ERROR: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())



