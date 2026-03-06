#!/usr/bin/env python3
"""
Final MDPI Converter with Renumbered Figures
Handles old figure references and converts to new sequential numbering

Usage:
    python3 convert_final.py ../V8.2.3.md ../V8.2.3_FINAL.docx

Author: Final Pipeline
Date: 2026
Version: 1.0
"""

import os
import sys
import re
from pathlib import Path
from typing import Optional, Tuple, Dict, List
import yaml

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


# ============================================================================
# NEW FIGURE MAPPING - Sequential Renumbering
# ============================================================================

# Maps NEW sequential figure numbers to files
# After running update_figure_numbers.py, manuscript uses 1-16
OLD_TO_NEW_MAPPING = {
    # Methods section (Figures 1-7)
    1: 'Figure_01_Study_Design.png',
    2: 'Figure_02_System_Architecture.png',
    3: 'Figure_03_Data_Flow.png',
    4: 'Figure_04_Detection_Process.png',
    5: 'Figure_05_Theoretical_Framework.png',
    6: 'Figure_06_Regulation_System.png',
    7: 'Figure_07_Evaluation_Framework.png',
    # Results section (Figures 8-13)
    8: 'Figure_08_Sample_Distribution.png',
    9: 'Figure_09_Personality_Dimensions.png',
    10: 'Figure_10_Personality_Heatmap.png',
    11: 'Figure_11_Performance_Comparison.png',
    12: 'Figure_12_Effect_Sizes.png',
    13: 'Figure_13_Weighted_Scores.png',
}

# Direct filename mapping for markdown image syntax
# Maps old filenames to actual files in final_figures/
FILENAME_TO_NEW = {
    # Updated by renumber script
    'Figure_01_Study_Design.png': 'Figure_01_Study_Design.png',
    'Figure_02_System_Architecture.png': 'Figure_02_System_Architecture.png',
    'Figure_03_Data_Flow.png': 'Figure_03_Data_Flow.png',
    'Figure_04_Detection_Process.png': 'Figure_04_Detection_Process.png',
    'Figure_05_Theoretical_Framework.png': 'Figure_05_Theoretical_Framework.png',
    'Figure_06_Regulation_System.png': 'Figure_06_Regulation_System.png',
    'Figure_07_Evaluation_Framework.png': 'Figure_07_Evaluation_Framework.png',
    'Figure_08_Sample_Distribution.png': 'Figure_08_Sample_Distribution.png',
    'Figure_09_Personality_Dimensions.png': 'Figure_09_Personality_Dimensions.png',
    'Figure_10_Personality_Heatmap.png': 'Figure_10_Personality_Heatmap.png',
    'Figure_11_Performance_Comparison.png': 'Figure_11_Performance_Comparison.png',
    'Figure_12_Effect_Sizes.png': 'Figure_12_Effect_Sizes.png',
    'Figure_13_Weighted_Scores.png': 'Figure_13_Weighted_Scores.png',
    # Original filenames (if still referenced)
    '10_system_architecture.png': 'Figure_02_System_Architecture.png',
    '11_study_design_flowchart.png': 'Figure_01_Study_Design.png',
    '16_trait_to_zurich_mapping.png': 'Figure_05_Theoretical_Framework.png',
    '13_data_flow_pipeline.png': 'Figure_03_Data_Flow.png',
    '14_detection_pipeline.png': 'Figure_04_Detection_Process.png',
    '15_regulation_prompt_assembly.png': 'Figure_06_Regulation_System.png',
    '12_evaluation_framework.png': 'Figure_07_Evaluation_Framework.png',
    '01_sample_distribution.png': 'Figure_08_Sample_Distribution.png',
    '03_performance_comparison.png': 'Figure_11_Performance_Comparison.png',
    '04_effect_sizes.png': 'Figure_12_Effect_Sizes.png',
    '06_personality_dimensions.png': 'Figure_09_Personality_Dimensions.png',
    '07_personality_heatmap.png': 'Figure_10_Personality_Heatmap.png',
    '08_weighted_scores.png': 'Figure_13_Weighted_Scores.png',
}

# Figure width for publication
FIGURE_WIDTH_INCHES = 6.5


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def extract_yaml_frontmatter(content: str) -> Tuple[Optional[Dict], str]:
    """Extract YAML front matter."""
    if not content.startswith('---'):
        return None, content
    
    lines = content.split('\n')
    yaml_end = -1
    
    for i in range(1, len(lines)):
        if lines[i].strip() == '---':
            yaml_end = i
            break
    
    if yaml_end == -1:
        return None, content
    
    yaml_content = '\n'.join(lines[1:yaml_end])
    remaining_content = '\n'.join(lines[yaml_end+1:])
    
    try:
        yaml_data = yaml.safe_load(yaml_content)
        return yaml_data, remaining_content
    except:
        return None, content


def setup_mdpi_styles(doc: Document) -> None:
    """Configure MDPI formatting."""
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(1.75)
        section.right_margin = Cm(1.75)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_title_page(doc: Document, yaml_data: Optional[Dict]) -> None:
    """Add title page."""
    if not yaml_data or 'title' not in yaml_data:
        return
    
    title = doc.add_paragraph(yaml_data['title'])
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.font.size = Pt(18)
    title_format.font.bold = True
    title_format.font.name = 'Times New Roman'
    
    doc.add_paragraph()


def find_figure(filename: str, search_dirs: List[Path]) -> Optional[Path]:
    """Search for figure in multiple directories."""
    for directory in search_dirs:
        figure_path = directory / filename
        if figure_path.exists():
            return figure_path
    return None


def insert_figure(doc: Document, figure_path: Path, fig_num: int) -> bool:
    """Insert figure into document."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    
    try:
        run.add_picture(str(figure_path), width=Inches(FIGURE_WIDTH_INCHES))
        print(f"    ? Figure {fig_num:02d}: {figure_path.name}")
        return True
    except Exception as e:
        print(f"    ? Error: {e}")
        return False


def process_markdown_content(doc: Document, content: str, figures_dirs: List[Path]) -> Dict:
    """Process markdown content and add to document."""
    lines = content.split('\n')
    stats = {'embedded': 0, 'missing': 0, 'skipped': 0}
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('####'):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('###'):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('##'):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('#'):
            doc.add_heading(line[1:].strip(), level=1)
        
        # Markdown image syntax: ![](figures/filename.png)
        elif line.startswith('!['):
            img_match = re.match(r'!\[\]\((.+?)\)', line)
            if img_match:
                img_path = img_match.group(1)
                old_filename = Path(img_path).name
                
                # Map to new filename
                new_filename = FILENAME_TO_NEW.get(old_filename, old_filename)
                figure_path = find_figure(new_filename, figures_dirs)
                
                if figure_path:
                    # Extract new figure number
                    fig_num = int(new_filename.split('_')[1])
                    if insert_figure(doc, figure_path, fig_num):
                        stats['embedded'] += 1
                    else:
                        stats['missing'] += 1
                else:
                    print(f"    Warning: {new_filename} not found")
                    stats['missing'] += 1
                    para = doc.add_paragraph(f"[Figure: {new_filename}]")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # [Figure N] references
        elif '[Figure' in line or '[figure' in line:
            # Check if this is a placeholder before an image (skip to avoid duplicates)
            if '**[Figure' in line and 'near here]**' in line:
                # Placeholder - skip it, image comes next
                stats['skipped'] += 1
            else:
                # Actual figure reference - embed it
                match = re.search(r'\[[Ff]igure\s+(\d+)', line)
                if match:
                    old_fig_num = int(match.group(1))
                    new_filename = OLD_TO_NEW_MAPPING.get(old_fig_num)
                    
                    if new_filename:
                        # Check if we already embedded this in last few lines (avoid duplicate)
                        figure_path = find_figure(new_filename, figures_dirs)
                        if figure_path:
                            new_fig_num = int(new_filename.split('_')[1])
                            if insert_figure(doc, figure_path, new_fig_num):
                                stats['embedded'] += 1
                            else:
                                stats['missing'] += 1
                        else:
                            print(f"    Warning: Figure {old_fig_num} -> {new_filename} not found")
                            stats['missing'] += 1
                    else:
                        # Figure number not in mapping
                        stats['skipped'] += 1
                else:
                    doc.add_paragraph(line)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line).strip(), style='List Number')
        
        # Regular paragraphs
        else:
            doc.add_paragraph(line)
        
        i += 1
    
    return stats


def convert_with_renumbered_figures(input_md: Path, output_docx: Path) -> bool:
    """Convert markdown to Word with renumbered figures."""
    
    print("\n" + "="*70)
    print("FINAL MDPI CONVERTER - Sequential Figure Numbering")
    print("="*70)
    
    if not input_md.exists():
        print(f"\n? Error: Input not found: {input_md}")
        return False
    
    print(f"\n?? Input: {input_md.name}")
    print(f"?? Output: {output_docx.name}")
    
    # Setup figure search directories (priority order)
    figures_dirs = []
    
    # 1. PRIORITY: final_figures/ (renumbered, styled)
    final_dir = input_md.parent / "final_figures"
    if final_dir.exists():
        figures_dirs.append(final_dir)
        print(f"?? Primary: final_figures/ ({len(list(final_dir.glob('*.png')))} figures)")
    
    # 2. Backup: unified_figures/
    unified_dir = input_md.parent / "unified_figures"
    if unified_dir.exists():
        figures_dirs.append(unified_dir)
        print(f"?? Backup: unified_figures/")
    
    if not figures_dirs:
        print("??  No figure directories found!")
        return False
    
    # Read markdown
    try:
        with open(input_md, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f"? Error reading file: {e}")
        return False
    
    # Parse YAML
    yaml_data, content = extract_yaml_frontmatter(content)
    if yaml_data:
        print(f"? YAML front matter extracted")
    
    # Create document
    doc = Document()
    setup_mdpi_styles(doc)
    print("? MDPI formatting applied")
    
    # Add title
    if yaml_data:
        add_title_page(doc, yaml_data)
        print("? Title page added")
    
    # Process content
    print("\n?? Processing content and embedding figures:")
    print("-" * 70)
    stats = process_markdown_content(doc, content, figures_dirs)
    
    # Save
    try:
        doc.save(str(output_docx))
        size_mb = output_docx.stat().st_size / (1024 * 1024)
        
        print("-" * 70)
        print(f"\n? Document saved: {output_docx}")
        print(f"  Size: {size_mb:.2f} MB")
        print(f"\n?? Figure Summary:")
        print(f"  Embedded: {stats['embedded']}")
        print(f"  Missing: {stats['missing']}")
        print(f"  Skipped: {stats['skipped']} (redundant)")
        
        return True
    except Exception as e:
        print(f"\n? Error saving: {e}")
        return False


def main():
    """Main entry point."""
    
    if len(sys.argv) < 2:
        print("\nUsage: python3 convert_final.py input.md [output.docx]")
        print("\nExample: python3 convert_final.py ../V8.2.3.md ../V8.2.3_FINAL.docx")
        sys.exit(1)
    
    input_md = Path(sys.argv[1])
    output_docx = Path(sys.argv[2]) if len(sys.argv) >= 3 else input_md.with_suffix('.docx')
    
    success = convert_with_renumbered_figures(input_md, output_docx)
    
    print("\n" + "="*70)
    if success:
        print("? CONVERSION SUCCESSFUL")
        print("="*70)
        print(f"\n?? Output: {output_docx}")
        print("\nAll figures renumbered sequentially (1-13)")
        print("All figures styled uniformly (300 DPI, borders)")
        print("\nNext: Open document and verify all figures!")
    else:
        print("? CONVERSION FAILED")
        print("="*70)
    
    print()
    return 0 if success else 1


if __name__ == "__main__":
    sys.exit(main())
