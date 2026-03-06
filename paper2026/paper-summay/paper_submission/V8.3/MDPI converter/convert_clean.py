#!/usr/bin/env python3
"""
Clean MDPI Markdown to Word Converter
======================================
Streamlined converter for academic manuscripts with MDPI formatting.

Features:
- YAML front matter extraction
- Automatic figure insertion from statistical analysis pipeline
- MDPI-compliant formatting (Times New Roman, 1.5 spacing, A4)
- Clean, maintainable code

Usage:
    python3 convert_clean.py input.md output.docx

Author: Unified Pipeline
Date: 2026
Version: 1.0
"""

import os
import sys
import re
import shutil
from pathlib import Path
from typing import Optional, Tuple, Dict
import yaml

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed")
    print("Install with: pip install python-docx PyYAML")
    sys.exit(1)


# ============================================================================
# CONFIGURATION
# ============================================================================

# Figure mapping - includes both old and new numbering
FIGURE_MAPPING = {
    # New numbering (reorganized V2.0)
    1: '01_performance_comparison.png',
    2: '02_effect_sizes.png',
    3: '03_personality_needs.png',
    4: '04_sample_quality.png',
    5: '05_personality_profiles.png',
    6: '06_system_architecture.png',
    7: '07_study_workflow.png',
    # Old numbering (for backward compatibility with V8.2.3)
    8: '08_weighted_scores.png',
    9: '09_total_score_boxplot.png',
    10: '10_system_overview.png',
    11: '11_study_workflow.png',
    12: '06_personality_dimensions.png',
    13: '01_sample_distribution.png',
    14: '03_performance_comparison.png',
    15: '04_effect_sizes.png',
    16: '07_personality_heatmap.png'
}

# MDPI formatting standards
MDPI_CONFIG = {
    'font_name': 'Times New Roman',
    'font_size': 12,
    'line_spacing': 1.5,
    'margin_top': 1.5,      # cm
    'margin_bottom': 3.5,   # cm
    'margin_left': 1.75,    # cm
    'margin_right': 1.75,   # cm
    'figure_width': 6.0,    # inches
}


# ============================================================================
# YAML FRONT MATTER
# ============================================================================

def extract_yaml_frontmatter(content: str) -> Tuple[Optional[Dict], str]:
    """
    Extract YAML front matter from markdown content.
    
    Returns:
        Tuple of (yaml_data dict or None, remaining content)
    """
    if not content.startswith('---'):
        return None, content
    
    lines = content.split('\n')
    yaml_end = -1
    
    # Find end of YAML block
    for i in range(1, len(lines)):
        if lines[i].strip() == '---':
            yaml_end = i
            break
    
    if yaml_end == -1:
        return None, content
    
    # Parse YAML
    yaml_content = '\n'.join(lines[1:yaml_end])
    remaining_content = '\n'.join(lines[yaml_end+1:])
    
    try:
        yaml_data = yaml.safe_load(yaml_content)
        return yaml_data, remaining_content
    except Exception as e:
        print(f"Warning: Could not parse YAML: {e}")
        return None, content


# ============================================================================
# DOCUMENT SETUP
# ============================================================================

def setup_mdpi_styles(doc: Document) -> None:
    """Configure document with MDPI formatting standards."""
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(MDPI_CONFIG['margin_top'])
        section.bottom_margin = Cm(MDPI_CONFIG['margin_bottom'])
        section.left_margin = Cm(MDPI_CONFIG['margin_left'])
        section.right_margin = Cm(MDPI_CONFIG['margin_right'])
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)
    
    # Configure Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = MDPI_CONFIG['font_name']
    normal_style.font.size = Pt(MDPI_CONFIG['font_size'])
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = MDPI_CONFIG['font_name']
        heading_style.font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        
        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)


def add_title_page(doc: Document, yaml_data: Optional[Dict]) -> None:
    """Add title page from YAML front matter."""
    if not yaml_data or 'title' not in yaml_data:
        return
    
    # Title
    title = doc.add_paragraph(yaml_data['title'])
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_format = title.runs[0]
    title_format.font.size = Pt(18)
    title_format.font.bold = True
    title_format.font.name = MDPI_CONFIG['font_name']
    
    # Add spacing
    doc.add_paragraph()


# ============================================================================
# CONTENT PROCESSING
# ============================================================================

def process_markdown_content(doc: Document, content: str, figures_dirs: list) -> None:
    """
    Process markdown content and add to document.
    
    Handles:
    - Headings (##, ###, ####)
    - Bold and italic
    - Figure references
    - Citations
    - Lists
    - Regular paragraphs
    
    Args:
        doc: Document object
        content: Markdown content
        figures_dirs: List of directories to search for figures
    """
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('####'):
            text = line[4:].strip()
            para = doc.add_heading(text, level=3)
        elif line.startswith('###'):
            text = line[3:].strip()
            para = doc.add_heading(text, level=2)
        elif line.startswith('##'):
            text = line[2:].strip()
            para = doc.add_heading(text, level=1)
        elif line.startswith('#'):
            text = line[1:].strip()
            para = doc.add_heading(text, level=1)
        
        # Figure references - both [Figure N] and ![](path) formats
        elif '[Figure' in line or '[figure' in line or line.startswith('!['):
            # Handle markdown image syntax: ![](figures/filename.png)
            img_match = re.match(r'!\[\]\((.+?)\)', line)
            if img_match:
                img_path = img_match.group(1)
                # Extract filename
                filename = Path(img_path).name
                # Try to find in figure directories
                figure_path = find_figure_in_directories(filename, figures_dirs)
                if figure_path:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    try:
                        run.add_picture(str(figure_path), width=Inches(MDPI_CONFIG['figure_width']))
                        print(f"  ✓ Inserted: {filename} (from {figure_path.parent.name})")
                    except Exception as e:
                        print(f"  ✗ Error inserting {filename}: {e}")
                        para = doc.add_paragraph(f"[Image: {filename}]")
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    print(f"  ⚠ Warning: Image not found: {filename}")
                    para = doc.add_paragraph(f"[Image: {filename}]")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Extract figure number from [Figure N] format
                match = re.search(r'\[[Ff]igure\s+(\d+)', line)
                if match:
                    fig_num = int(match.group(1))
                    insert_figure(doc, fig_num, figures_dirs)
                else:
                    para = doc.add_paragraph(line)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            para = doc.add_paragraph(text, style='List Bullet')
            para = apply_text_formatting(para, text)
        
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            para = doc.add_paragraph(text, style='List Number')
            para = apply_text_formatting(para, text)
        
        # Regular paragraph
        else:
            para = doc.add_paragraph(line)
            para = apply_text_formatting(para, line)
        
        i += 1


def apply_text_formatting(para, text: str):
    """Apply bold, italic, and other inline formatting."""
    para.clear()
    
    # Process bold (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)
    
    return para


# ============================================================================
# FIGURE INSERTION
# ============================================================================

def find_figure_in_directories(filename: str, search_dirs: list) -> Optional[Path]:
    """
    Search for a figure file in multiple directories.
    
    Args:
        filename: Figure filename to search for
        search_dirs: List of directories to search
        
    Returns:
        Path to figure if found, None otherwise
    """
    for directory in search_dirs:
        figure_path = directory / filename
        if figure_path.exists():
            return figure_path
    return None


def insert_figure(doc: Document, fig_num: int, figures_dirs: list) -> bool:
    """
    Insert a figure into the document.
    Searches multiple directories for the figure.
    
    Args:
        doc: Document object
        fig_num: Figure number
        figures_dirs: List of Path objects to search for figures
        
    Returns:
        True if figure inserted successfully, False otherwise
    """
    filename = FIGURE_MAPPING.get(fig_num)
    if not filename:
        print(f"  Warning: Unknown figure number: {fig_num}")
        return False
    
    # Search for figure in all directories
    figure_path = find_figure_in_directories(filename, figures_dirs)
    
    if not figure_path:
        print(f"  Warning: Figure not found: {filename}")
        # Add placeholder
        para = doc.add_paragraph(f"[Figure {fig_num}: {filename}]")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.runs[0].italic = True
        return False
    
    # Insert figure
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    
    try:
        run.add_picture(str(figure_path), width=Inches(MDPI_CONFIG['figure_width']))
        print(f"  ✓ Inserted: Figure {fig_num} - {filename} (from {figure_path.parent.name})")
        return True
    except Exception as e:
        print(f"  ✗ Error inserting figure {fig_num}: {e}")
        return False


# ============================================================================
# MAIN CONVERSION
# ============================================================================

def convert_markdown_to_docx(input_md: Path, output_docx: Path, 
                            figures_dir: Optional[Path] = None) -> bool:
    """
    Convert markdown file to MDPI-formatted Word document.
    
    Args:
        input_md: Path to input markdown file
        output_docx: Path to output Word document
        figures_dir: Path to figures directory (auto-detected if None)
        
    Returns:
        True if conversion successful, False otherwise
    """
    print("\n" + "="*70)
    print("MDPI MARKDOWN TO WORD CONVERTER")
    print("="*70)
    
    # Validate input
    if not input_md.exists():
        print(f"✗ Error: Input file not found: {input_md}")
        return False
    
    print(f"\n📄 Input: {input_md.name}")
    print(f"📄 Output: {output_docx.name}")
    
    # Build list of figure directories to search
    figures_dirs = []
    
    if figures_dir is not None:
        # User-specified directory
        if figures_dir.exists():
            figures_dirs.append(figures_dir)
            print(f"📁 Figure search: {figures_dir}")
    else:
        # Auto-detect figure directories (priority order)
        
        # 1. PRIORITY: Unified figures directory (consolidated, standardized)
        unified_dir = input_md.parent / "unified_figures"
        if unified_dir.exists():
            figures_dirs.append(unified_dir)
            print(f"📁 Figure search: unified_figures/ (primary)")
        
        # 2. Statistical analysis figures (generated)
        stat_analysis_dir = input_md.parent / "statistical analyis" / "figures"
        if stat_analysis_dir.exists():
            figures_dirs.append(stat_analysis_dir)
            print(f"📁 Figure search: statistical analyis/figures/ (backup)")
        
        # 3. Main figures directory (architectural diagrams)
        main_figures_dir = input_md.parent.parent / "figures"
        if main_figures_dir.exists():
            figures_dirs.append(main_figures_dir)
            print(f"📁 Figure search: ../figures/ (backup)")
        
        # 4. V8.3 figures directory
        v83_figures_dir = input_md.parent / "figures"
        if v83_figures_dir.exists():
            figures_dirs.append(v83_figures_dir)
            print(f"📁 Figure search: V8.3/figures/ (backup)")
    
    if not figures_dirs:
        print("⚠️  Warning: No figure directories found")
        figures_dirs = [input_md.parent]  # Fallback to document directory
    
    # Read markdown content
    try:
        with open(input_md, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        print(f"✗ Error reading file: {e}")
        return False
    
    # Extract YAML front matter
    yaml_data, content = extract_yaml_frontmatter(content)
    if yaml_data:
        print(f"✓ Extracted YAML front matter")
        if 'title' in yaml_data:
            print(f"  Title: {yaml_data['title'][:60]}...")
    
    # Create document
    doc = Document()
    setup_mdpi_styles(doc)
    print("✓ Applied MDPI formatting")
    
    # Add title page
    if yaml_data:
        add_title_page(doc, yaml_data)
        print("✓ Added title page")
    
    # Process content
    print("\n📝 Processing content...")
    print(f"   Searching in {len(figures_dirs)} figure director{'y' if len(figures_dirs) == 1 else 'ies'}")
    process_markdown_content(doc, content, figures_dirs)
    
    # Save document
    try:
        doc.save(str(output_docx))
        print(f"\n✓ Document saved: {output_docx}")
        
        # File size
        size_mb = output_docx.stat().st_size / (1024 * 1024)
        print(f"  Size: {size_mb:.2f} MB")
        
        return True
    except Exception as e:
        print(f"✗ Error saving document: {e}")
        return False


# ============================================================================
# COMMAND LINE INTERFACE
# ============================================================================

def main():
    """Main entry point."""
    print("\n" + "="*70)
    print("CLEAN MDPI CONVERTER V1.0")
    print("="*70)
    
    # Parse arguments
    if len(sys.argv) < 2:
        print("\nUsage: python3 convert_clean.py input.md [output.docx] [figures_dir]")
        print("\nExamples:")
        print("  python3 convert_clean.py V8.2.3.md")
        print("  python3 convert_clean.py V8.2.3.md V8.2.3_output.docx")
        print("  python3 convert_clean.py V8.2.3.md output.docx ../figures")
        print()
        sys.exit(1)
    
    # Input file
    input_md = Path(sys.argv[1])
    
    # Output file (default: same name with .docx extension)
    if len(sys.argv) >= 3:
        output_docx = Path(sys.argv[2])
    else:
        output_docx = input_md.with_suffix('.docx')
    
    # Figures directory (optional)
    figures_dir = None
    if len(sys.argv) >= 4:
        figures_dir = Path(sys.argv[3])
    
    # Convert
    success = convert_markdown_to_docx(input_md, output_docx, figures_dir)
    
    # Summary
    print("\n" + "="*70)
    if success:
        print("? CONVERSION SUCCESSFUL")
        print("="*70)
        print(f"\n?? Output: {output_docx}")
        print("\nNext steps:")
        print("  1. Open document in Word")
        print("  2. Verify formatting and figures")
        print("  3. Review for MDPI compliance")
    else:
        print("? CONVERSION FAILED")
        print("="*70)
        print("\nPlease check the error messages above.")
    
    print()
    return 0 if success else 1


if __name__ == "__main__":
    sys.exit(main())
