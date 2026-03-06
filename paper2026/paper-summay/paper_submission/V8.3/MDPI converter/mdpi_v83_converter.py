#!/usr/bin/env python3
"""
Enhanced MDPI Markdown to Word Converter - V8.3 Specific
-------------------------------------------------------
Converts Markdown manuscripts to MDPI-compliant Word documents.
Updated for V8.3 structure with dual figure sources (MDPI diagrams + Statistical plots).
"""

import os
import sys
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MDPIConverterV83:
    """Enhanced MDPI manuscript converter for V8.3."""
    
    def __init__(self, md_file, output_dir):
        """Initialize converter."""
        self.md_file = Path(md_file)
        self.base_dir = self.md_file.parent
        self.output_dir = Path(output_dir)
        self.doc = None
        self.table_counter = 1
        
        # Create output directory
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Define figure sources
        self.figures_dir = self.base_dir / "figures"
        self.mdpi_figs = self.figures_dir / "mdpi"
        
        # Figure mapping for V8.3
        self.figure_mapping = {
            1: (self.mdpi_figs, 'study_design_mdpi.png'),
            2: (self.mdpi_figs, 'system_architecture_mdpi.png'),
            3: (self.mdpi_figs, 'data_flow_mdpi.png'),
            4: (self.mdpi_figs, 'detection_pipeline_mdpi.png'),
            5: (self.mdpi_figs, 'trait_mapping_mdpi.png'),
            6: (self.mdpi_figs, 'regulation_workflow_mdpi.png'),
            7: (self.mdpi_figs, 'evaluation_framework_mdpi.png'),
            8: (self.figures_dir, 'data_quality_summary.png'),
            9: (self.figures_dir, '06_personality_dimensions.png'),
            10: (self.figures_dir, '07_personality_heatmap.png'),
            11: (self.figures_dir, '08_weighted_scores.png'),
            12: (self.figures_dir, '09_total_score_boxplot.png'),
            13: (self.figures_dir, '10_selective_enhancement_paired.png'),
            14: (self.figures_dir, '11_metric_composition.png'),
            15: (self.figures_dir, 'dialogue_illustration_1.png'),
            16: (self.figures_dir, 'dialogue_illustration_2.png')
        }

    def setup_mdpi_formatting(self):
        """Set up MDPI margins and basic font."""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(1.75)
            section.right_margin = Cm(1.75)
        
        self.create_mdpi_styles()

    def create_mdpi_styles(self):
        """Create MDPI styles."""
        styles = self.doc.styles
        
        # Normal Text
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Figure Caption
        try:
            fig_cap = styles.add_style('MDPI Figure Caption', WD_STYLE_TYPE.PARAGRAPH)
            fig_cap.font.name = 'Times New Roman'
            fig_cap.font.size = Pt(10)
            fig_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_cap.paragraph_format.space_after = Pt(12)
        except ValueError: pass

        # Table Caption
        try:
            tab_cap = styles.add_style('MDPI Table Caption', WD_STYLE_TYPE.PARAGRAPH)
            tab_cap.font.name = 'Times New Roman'
            tab_cap.font.size = Pt(10)
            tab_cap.font.bold = True
            tab_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except ValueError: pass

    def insert_figure(self, fig_num, caption_text):
        """Insert figure into document."""
        if fig_num not in self.figure_mapping:
            print(f"Warning: No mapping for Figure {fig_num}")
            return False
        
        dir_path, filename = self.figure_mapping[fig_num]
        fig_path = dir_path / filename
        
        if not fig_path.exists():
            print(f"Warning: Figure file not found: {fig_path}")
            return False
            
        # Add Image
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(fig_path), width=Inches(6.0))
        
        # Add Caption
        cap_para = self.doc.add_paragraph()
        cap_para.style = 'MDPI Figure Caption'
        
        # Bold "Figure X."
        bold_run = cap_para.add_run(f"Figure {fig_num}. ")
        bold_run.bold = True
        
        # Rest of caption
        clean_cap = re.sub(r'^\*\*Figure\s+\d+\.\*\*\s*', '', caption_text)
        cap_para.add_run(clean_cap)
        
        print(f"? Inserted Figure {fig_num}")
        return True

    def add_mdpi_table(self, table_lines, caption=None):
        """Add MDPI formatted table."""
        rows = []
        for line in table_lines:
            if re.match(r'^\|[\s\-\:]+\|', line): continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells: rows.append(cells)
            
        if not rows: return

        if caption:
            cap_para = self.doc.add_paragraph()
            cap_para.style = 'MDPI Table Caption'
            bold_run = cap_para.add_run(f"Table {self.table_counter}. ")
            bold_run.bold = True
            cap_para.add_run(caption)
            self.table_counter += 1

        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = re.sub(r'\*\*|\*|`', '', cell_data)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(10)
                        if i == 0: r.bold = True
        
        self.doc.add_paragraph()

    def add_markdown_runs(self, paragraph, text: str):
        """Add runs to a paragraph, handling basic Markdown emphasis."""
        if not text:
            return

        # Normalize common LaTeX/math markers into plain text
        # (MDPI Word submission does not require LaTeX rendering here.)
        replacements = {
            r"\times": "×",
            r"\cdot": "·",
            r"\pm": "±",
            r"\leq": "≤",
            r"\geq": "≥",
            r"\rightarrow": "→",
            r"\to": "→",
            r"\alpha": "α",
            r"\beta": "β",
            r"\gamma": "γ",
            r"\delta": "δ",
            r"\kappa": "κ",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)

        # Remove simple inline LaTeX wrappers and math delimiters
        text = text.replace(r"\(", "").replace(r"\)", "")
        text = text.replace(r"\[", "").replace(r"\]", "")
        text = text.replace("$", "")

        # Unescape common LaTeX escapes used in Markdown drafts
        text = text.replace(r"\_", "_").replace(r"\%", "%")

        bold = False
        italic = False
        buf = []
        i = 0

        def flush():
            if not buf:
                return
            run = paragraph.add_run("".join(buf))
            run.bold = bold
            run.italic = italic
            buf.clear()

        while i < len(text):
            # Bold toggle: **
            if text.startswith("**", i):
                flush()
                bold = not bold
                i += 2
                continue

            # Italic toggle: * (only if not part of **)
            if text[i] == "*":
                flush()
                italic = not italic
                i += 1
                continue

            # Inline code: `code`
            if text[i] == "`":
                flush()
                j = text.find("`", i + 1)
                if j == -1:
                    i += 1
                    continue
                code_text = text[i + 1 : j]
                run = paragraph.add_run(code_text)
                run.font.name = "Courier New"
                i = j + 1
                continue

            buf.append(text[i])
            i += 1

        flush()

    def convert(self, output_name):
        """Main conversion loop."""
        self.doc = Document()
        self.setup_mdpi_formatting()
        
        with open(self.md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        i = 0
        skip_mode = False
        title = ""
        
        # First pass: Extract Title from YAML and handle YAML skipping
        if i < len(lines) and lines[i].strip() == '---':
            i += 1
            while i < len(lines) and lines[i].strip() != '---':
                line = lines[i].strip()
                if line.startswith('title:'):
                    title = line.replace('title:', '').strip().strip('"').strip("'")
                i += 1
            i += 1 # Skip the closing ---
        
        # Insert Title if found
        if title:
            h = self.doc.add_heading(title, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"✓ Added Title: {title[:50]}...")

        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue

            # Detect and skip Changelog/Notes section
            if line.startswith('##') and ('Changelog' in line or 'editorial' in line):
                skip_mode = True
                i += 1
                continue
            
            # End skip mode when a new major section starts or Abstract
            if skip_mode and (line.startswith('# ') or line.startswith('## Abstract')):
                skip_mode = False
            
            if skip_mode:
                i += 1
                continue
                
            # Headings
            if line.startswith('#'):
                level = line.count('#')
                text = line.replace('#', '').strip()
                self.doc.add_heading(text, level=min(level, 3))
                i += 1
                continue

            # Figure Placeholders
            fig_match = re.match(
                r'^(?:\*\*)?\[Figure\s+(\d+)\s+near\s+here\](?:\*\*)?$',
                line,
                re.IGNORECASE
            )
            if fig_match:
                fig_num = int(fig_match.group(1))
                # Find caption
                caption = ""
                j = i + 1
                while j < len(lines):
                    if lines[j].strip().startswith('**Figure'):
                        caption = lines[j].strip()
                        i = j
                        break
                    j += 1
                self.insert_figure(fig_num, caption)
                i += 1
                continue

            def _is_table_separator(s: str) -> bool:
                s = s.strip()
                # typical Markdown separator line: |----|:---:|----|
                return bool(re.match(r'^\|\s*[:\-]+\s*(\|\s*[:\-]+\s*)+\|?\s*$', s))

            def _is_table_row(s: str) -> bool:
                s = s.strip()
                return s.startswith('|') and '|' in s[1:]

            # Markdown table (caption may be before or after)
            if _is_table_row(line):
                # Confirm it's a real table by checking the next non-empty line is a separator
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1
                if j < len(lines) and _is_table_separator(lines[j]):
                    table_lines = []
                    while i < len(lines) and _is_table_row(lines[i]):
                        table_lines.append(lines[i])
                        i += 1

                    # Optional caption AFTER the table (common in this manuscript)
                    caption = None
                    k = i
                    while k < len(lines) and not lines[k].strip():
                        k += 1
                    if k < len(lines):
                        after = lines[k].strip()
                        if after.startswith('**Table'):
                            caption = re.sub(r'^\*\*Table\s+\d+\.\*\*\s*', '', after).replace('**', '')
                            i = k + 1

                    self.add_mdpi_table(table_lines, caption=caption)
                    continue

            # Table
            if line.startswith('**Table'):
                caption = re.sub(r'^\*\*Table\s+\d+\.\*\*\s*', '', line).replace('**', '')
                table_lines = []
                i += 1
                # Skip blank lines between caption and the table
                while i < len(lines) and not lines[i].strip():
                    i += 1
                while i < len(lines) and _is_table_row(lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                self.add_mdpi_table(table_lines, caption)
                continue

            # Paragraphs
            if line:
                # Bullets / numbering
                bullet_match = re.match(r'^-\s+(.*)$', line)
                num_match = re.match(r'^\d+\.\s+(.*)$', line)
                if bullet_match:
                    p = self.doc.add_paragraph(style='List Bullet')
                    text = bullet_match.group(1).strip()
                elif num_match:
                    p = self.doc.add_paragraph(style='List Number')
                    text = num_match.group(1).strip()
                else:
                    p = self.doc.add_paragraph()
                    text = line

                # Ensure refs are simple
                text = re.sub(r'\[(\d+)\]', r'[\1]', text)

                # Basic Markdown emphasis
                self.add_markdown_runs(p, text)
                
            i += 1
            
        out_path = self.output_dir / output_name
        self.doc.save(str(out_path))
        print(f"\nFinal DOCX saved to: {out_path}")


if __name__ == "__main__":
    conv = MDPIConverterV83(
        "/Users/huaduojiejia/MyProject/hslu/2026/paper2026/paper-summay/paper_submission/V8.3/V8.2.7.md",
        "/Users/huaduojiejia/MyProject/hslu/2026/paper2026/paper-summay/paper_submission/V8.3/docoutput"
    )
    conv.convert("V8.2.7_Healthcare_Submission_MDPI.docx")
