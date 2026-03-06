#!/usr/bin/env python3
"""
MDPI Markdown to Word Converter
--------------------------------
Converts Markdown files to Microsoft Word (.docx) documents 
formatted according to MDPI (Multidisciplinary Digital Publishing Institute) guidelines.

MDPI Requirements:
- 12pt Times New Roman font
- 1.5 line spacing
- A4 paper size
- Standard margins (Top/Bottom: 1.5cm, Left/Right: 1.75cm)
- Numbered references [1], [2], etc.
- IMRaD structure (Introduction, Methods, Results, Discussion)
- Mandatory sections: Abstract, Keywords, Author Contributions, 
  Funding, Conflicts of Interest, Data Availability, IRB Statement
"""

import os
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_mdpi_formatting(doc):
    """Set up MDPI-specific document formatting."""
    # Set document margins according to MDPI standards
    # A4 paper with margins: Top/Bottom ~1.5cm, Left/Right ~1.75cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.59)    # ~1.5 cm
        section.bottom_margin = Inches(1.38)  # ~3.5 cm (more space for page numbers)
        section.left_margin = Inches(0.69)   # ~1.75 cm
        section.right_margin = Inches(0.69)  # ~1.75 cm
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
    
    # Create MDPI-specific styles
    create_mdpi_styles(doc)


def create_mdpi_styles(doc):
    """Create custom styles for MDPI formatting."""
    styles = doc.styles
    
    # MDPI Body Text style (12pt Times New Roman, 1.5 spacing)
    try:
        body_style = styles.add_style('MDPI Body', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(12)
        body_style.paragraph_format.line_spacing = 1.5
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except ValueError:
        # Style already exists
        body_style = styles['MDPI Body']
    
    # Update Normal style to match MDPI requirements
    try:
        normal_style = styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except KeyError:
        pass
    
    # MDPI Abstract style
    try:
        abstract_style = styles.add_style('MDPI Abstract', WD_STYLE_TYPE.PARAGRAPH)
        abstract_style.font.name = 'Times New Roman'
        abstract_style.font.size = Pt(12)
        abstract_style.paragraph_format.line_spacing = 1.5
        abstract_style.paragraph_format.space_after = Pt(6)
        abstract_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except ValueError:
        pass
    
    # MDPI Code style (for code blocks)
    try:
        code_style = styles.add_style('MDPI Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(10)
        code_style.paragraph_format.line_spacing = 1.0
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.5)
    except ValueError:
        pass
    
    # Update heading styles to match MDPI
    for i in range(1, 4):
        try:
            heading_style = styles[f'Heading {i}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True
            if i == 1:
                heading_style.font.size = Pt(16)
            elif i == 2:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass


def process_markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to MDPI-formatted Word document."""
    try:
        # Create new document
        doc = Document()
        
        # Apply MDPI formatting
        setup_mdpi_formatting(doc)
        
        # Read Markdown content
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Split content into lines
        lines = md_content.split('\n')
        
        # Track document structure
        in_code_block = False
        code_block_content = []
        in_table = False
        table_lines = []
        title_added = False
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_block_content = []
                else:
                    # End of code block
                    in_code_block = False
                    if code_block_content:
                        code_para = doc.add_paragraph('\n'.join(code_block_content))
                        try:
                            code_para.style = 'MDPI Code'
                        except KeyError:
                            code_para.style = 'Normal'
                            code_para.runs[0].font.name = 'Courier New'
                            code_para.runs[0].font.size = Pt(10)
                i += 1
                continue
            
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
            
            # Handle YAML front matter (skip it)
            if line.strip() == '---' and i == 0:
                # Skip until next ---
                i += 1
                while i < len(lines) and lines[i].strip() != '---':
                    i += 1
                i += 1
                continue
            
            # Handle headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                
                # Remove markdown formatting from headings
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                text = re.sub(r'\{.*?\}', '', text)  # Remove {.unnumbered} etc.
                text = text.strip()
                
                # Add heading to document
                if level == 1 and not title_added:
                    doc.add_heading(text, level=0)  # Title style
                    title_added = True
                elif level <= 3:
                    doc.add_heading(text, level=level)
                else:
                    # For h4-h6, use bold paragraph
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    run.bold = True
                    run.font.size = Pt(12)
                
                i += 1
                continue
            
            # Handle tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_lines = [line]
                else:
                    table_lines.append(line)
                i += 1
                continue
            elif in_table and not ('|' in line):
                # End of table
                in_table = False
                add_table_to_doc(doc, table_lines)
                table_lines = []
                # Don't increment i, process current line
                continue
            
            # Handle empty lines
            if not line.strip():
                i += 1
                continue
            
            # Handle bold/italic text and add as paragraph
            para_text = line.strip()
            
            # Skip horizontal rules
            if re.match(r'^[\-\*\_]{3,}$', para_text):
                i += 1
                continue
            
            # Add paragraph
            para = doc.add_paragraph()
            add_formatted_text(para, para_text)
            
            i += 1
        
        # Handle any remaining table
        if in_table and table_lines:
            add_table_to_doc(doc, table_lines)
        
        # Save document
        doc.save(docx_file)
        return True
        
    except Exception as e:
        print(f"Error converting document: {e}")
        import traceback
        traceback.print_exc()
        return False


def add_formatted_text(paragraph, text):
    """Add text with markdown formatting to a paragraph."""
    # Handle bold, italic, inline code
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold text
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        # Italic text
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        # Inline code
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        # Links [text](url)
        elif re.match(r'\[.*?\]\(.*?\)', part):
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if link_match:
                text_part = link_match.group(1)
                url = link_match.group(2)
                run = paragraph.add_run(text_part)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
        else:
            paragraph.add_run(part)


def add_table_to_doc(doc, table_lines):
    """Convert markdown table to Word table with MDPI formatting."""
    if not table_lines:
        return
    
    # Parse table
    rows = []
    for line in table_lines:
        # Skip separator lines (|---|---|)
        if re.match(r'^\|[\s\-\:]+\|', line):
            continue
        
        # Split by | and clean
        cells = [cell.strip() for cell in line.split('|')]
        cells = [c for c in cells if c]  # Remove empty cells
        
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create Word table
    num_cols = len(rows[0])
    num_rows = len(rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                # Remove markdown formatting
                cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_data)
                cell_text = re.sub(r'\*(.+?)\*', r'\1', cell_text)
                cell.text = cell_text
                
                # Format cell
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                    paragraph.paragraph_format.space_after = Pt(0)
                
                # Bold first row (header)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # Auto-fit table
    table.autofit = True
    
    # Add spacing after table
    doc.add_paragraph()


def main():
    """Main entry point."""
    if len(sys.argv) < 3:
        print("Usage: python mdpi_md2word.py <input.md> <output.docx>")
        sys.exit(1)
    
    md_file = sys.argv[1]
    docx_file = sys.argv[2]
    
    if not os.path.exists(md_file):
        print(f"Error: Input file '{md_file}' not found.")
        sys.exit(1)
    
    print(f"Converting {md_file} to MDPI format...")
    success = process_markdown_to_docx(md_file, docx_file)
    
    if success:
        print(f"Successfully converted to {docx_file}")
        sys.exit(0)
    else:
        print("Conversion failed.")
        sys.exit(1)


if __name__ == '__main__':
    main()




