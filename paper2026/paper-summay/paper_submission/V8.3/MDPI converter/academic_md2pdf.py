#!/usr/bin/env python3
"""
Enhanced Academic Markdown to PDF/DOCX Converter
------------------------------------------------
Converts Markdown files to professionally formatted academic documents
with improved styling for tables, code blocks, and overall presentation.

Features:
- PDF export via LibreOffice/Word conversion
- Enhanced table formatting with borders and shading
- Professional font selection (Times New Roman for body, Courier for code)
- Proper academic spacing and margins
- Automatic table of contents generation
- Section numbering
- Improved code block styling
- Better handling of references and appendices
"""

import os
import sys
import re
import subprocess
from pathlib import Path
import click
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class AcademicConverter:
    """Enhanced converter for academic documents with PDF support."""
    
    def __init__(self, template_file=None):
        self.template_file = template_file
        self.section_counter = {'h1': 0, 'h2': 0, 'h3': 0, 'h4': 0}
        
    def setup_styles(self, doc):
        """Set up professional academic styles for the document."""
        
        # Configure normal style
        try:
            normal_style = doc.styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(12)
            
            normal_paragraph = normal_style.paragraph_format
            normal_paragraph.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            normal_paragraph.space_after = Pt(6)
            normal_paragraph.first_line_indent = Inches(0.5)
        except KeyError:
            pass
        
        # Configure heading styles
        heading_configs = {
            'Heading 1': {'size': 16, 'bold': True, 'spacing_before': 12, 'spacing_after': 6},
            'Heading 2': {'size': 14, 'bold': True, 'spacing_before': 12, 'spacing_after': 6},
            'Heading 3': {'size': 12, 'bold': True, 'spacing_before': 6, 'spacing_after': 3},
            'Heading 4': {'size': 12, 'bold': True, 'italic': True, 'spacing_before': 6, 'spacing_after': 3},
        }
        
        for style_name, config in heading_configs.items():
            try:
                style = doc.styles[style_name]
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(config['size'])
                font.bold = config.get('bold', False)
                font.italic = config.get('italic', False)
                
                para_format = style.paragraph_format
                para_format.space_before = Pt(config.get('spacing_before', 0))
                para_format.space_after = Pt(config.get('spacing_after', 0))
                para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            except KeyError:
                continue
        
        # Create/configure code block style
        try:
            code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            code_style = doc.styles['Code Block']
        
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(10)
        
        code_para = code_style.paragraph_format
        code_para.left_indent = Inches(0.5)
        code_para.space_before = Pt(6)
        code_para.space_after = Pt(6)
        
        # Set up page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def add_table_of_contents(self, doc):
        """Add a professional table of contents."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # Add TOC field
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar3)
        
        # Add page break after TOC
        doc.add_page_break()
    
    def format_table(self, table):
        """Apply professional formatting to tables."""
        # Set table style
        table.style = 'Light Grid Accent 1'
        
        # Format header row
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # Set header background color (light gray)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9D9D9')
                cell._element.get_or_add_tcPr().append(shading_elm)
                
                # Bold header text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
        
        # Set all cells font
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(3)
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        if run.font.size is None or run.font.size > Pt(11):
                            run.font.size = Pt(10)
        
        # Auto-fit table to content
        table.autofit = True
    
    def convert_markdown_to_docx(self, md_file, docx_file, add_toc=True, number_sections=True):
        """
        Convert markdown to DOCX with enhanced academic formatting.
        
        Args:
            md_file: Path to input markdown file
            docx_file: Path to output DOCX file
            add_toc: Whether to add table of contents
            number_sections: Whether to add section numbering
        """
        # Create document
        if self.template_file and os.path.exists(self.template_file):
            doc = Document(self.template_file)
        else:
            doc = Document()
        
        # Set up styles
        self.setup_styles(doc)
        
        # Read markdown content
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Add TOC if requested
        if add_toc:
            doc.add_heading('Table of Contents', level=1)
            self.add_table_of_contents(doc)
        
        # Process content line by line
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []
        
        while i < len(lines):
            line = lines[i]
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                    lang = line.strip()[3:].strip()  # Get language identifier
                else:
                    # End of code block
                    in_code_block = False
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text, style='Code Block')
                    # Add subtle background
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.5)
                    code_lines = []
                i += 1
                continue
            
            if in_code_block:
                code_lines.append(line)
                i += 1
                continue
            
            # Handle tables
            if line.strip().startswith('|') and '|' in line:
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
                i += 1
                continue
            elif in_table and not line.strip().startswith('|'):
                # End of table
                in_table = False
                self.add_markdown_table(doc, table_lines)
                table_lines = []
                # Don't increment i, process current line
                continue
            
            # Handle headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                
                # Remove markdown formatting from headings
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                
                # Add section numbering if requested
                if number_sections and level <= 3:
                    section_num = self.get_section_number(level)
                    if section_num and not text.startswith(('Abstract', 'References', 'Appendix')):
                        text = "{} {}".format(section_num, text)
                
                doc.add_heading(text, level=min(level, 4))
                i += 1
                continue
            
            # Handle paragraphs
            if line.strip():
                p = doc.add_paragraph()
                self.add_formatted_text(p, line)
            
            i += 1
        
        # Save document
        doc.save(docx_file)
        print("[OK] DOCX created: {}".format(docx_file))
        return docx_file
    
    def get_section_number(self, level):
        """Generate section number based on current counters."""
        if level == 1:
            self.section_counter['h1'] += 1
            self.section_counter['h2'] = 0
            self.section_counter['h3'] = 0
            self.section_counter['h4'] = 0
            return str(self.section_counter['h1'])
        elif level == 2:
            self.section_counter['h2'] += 1
            self.section_counter['h3'] = 0
            self.section_counter['h4'] = 0
            return "{}.{}".format(self.section_counter['h1'], self.section_counter['h2'])
        elif level == 3:
            self.section_counter['h3'] += 1
            self.section_counter['h4'] = 0
            return "{}.{}.{}".format(self.section_counter['h1'], self.section_counter['h2'], self.section_counter['h3'])
        return None
    
    def add_formatted_text(self, paragraph, text):
        """Add text with markdown formatting (bold, italic, code)."""
        # Parse inline formatting
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                paragraph.add_run(part)
    
    def add_markdown_table(self, doc, table_lines):
        """Parse and add a markdown table with professional formatting."""
        if len(table_lines) < 2:
            return
        
        # Parse table structure
        rows = []
        for line in table_lines:
            if re.match(r'^\|[\s\-:|]+\|$', line):  # Skip separator line
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
        
        if not rows:
            return
        
        # Create table
        num_cols = len(rows[0])
        num_rows = len(rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        
        # Fill table
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
        
        # Apply formatting
        self.format_table(table)
        
        # Add spacing after table
        doc.add_paragraph()
    
    def convert_to_pdf(self, docx_file, pdf_file, method='libreoffice'):
        """
        Convert DOCX to PDF using available tools.
        
        Args:
            docx_file: Path to input DOCX file
            pdf_file: Path to output PDF file
            method: Conversion method ('libreoffice', 'word_mac', 'pandoc')
        """
        docx_path = Path(docx_file).resolve()
        pdf_path = Path(pdf_file).resolve()
        output_dir = pdf_path.parent
        
        if method == 'libreoffice':
            try:
                # Try LibreOffice conversion
                cmd = [
                    'soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(output_dir),
                    str(docx_path)
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    # LibreOffice creates PDF with same name as DOCX
                    generated_pdf = output_dir / "{docx_path.stem}.pd"
                    if generated_pdf.exists() and generated_pdf != pdf_path:
                        generated_pdf.rename(pdf_path)
                    print("[OK] PDF created: {}".format(pdf_file))
                    return True
                else:
                    print("[WARNING] LibreOffice conversion failed: {}".format(result.stderr))
                    return False
            except (subprocess.TimeoutExpired, FileNotFoundError) as e:
                print("[WARNING] LibreOffice not available: {}".format(e))
                return False
        
        elif method == 'word_mac':
            # macOS Word automation via AppleScript
            try:
                apple_script = f'''
                tell application "Microsoft Word"
                    open "{docx_path}"
                    save as active document file name "{pdf_path}" file format format PDF
                    close active document
                end tell
                '''
                result = subprocess.run(['osascript', '-e', apple_script], 
                                       capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    print("[OK] PDF created: {}".format(pdf_file))
                    return True
                else:
                    print("[WARNING] Word conversion failed: {}".format(result.stderr))
                    return False
            except Exception as e:
                print("[WARNING] Word automation failed: {}".format(e))
                return False
        
        elif method == 'pandoc':
            # Try pandoc with wkhtmltopdf
            try:
                cmd = [
                    'pandoc',
                    str(docx_path),
                    '-o', str(pdf_path),
                    '--pdf-engine=wkhtmltopdf'
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    print("[OK] PDF created: {}".format(pdf_file))
                    return True
                else:
                    print("[WARNING] Pandoc conversion failed: {}".format(result.stderr))
                    return False
            except Exception as e:
                print("[WARNING] Pandoc not available: {}".format(e))
                return False
        
        return False
    
    def auto_convert_to_pdf(self, docx_file, pdf_file):
        """Try multiple methods to convert DOCX to PDF."""
        methods = ['libreoffice', 'word_mac', 'pandoc']
        
        for method in methods:
            print("Attempting PDF conversion using {}...".format(method))
            if self.convert_to_pdf(docx_file, pdf_file, method=method):
                return True
        
        print("[WARNING] PDF conversion failed. Please convert manually from the DOCX file.")
        print("   DOCX file: {}".format(docx_file))
        return False


@click.command()
@click.argument('input_file', type=click.Path(exists=True))
@click.argument('output_file', type=click.Path(), required=False)
@click.option('--template', '-t', type=click.Path(exists=True), help='Word template file')
@click.option('--pdf', '-p', is_flag=True, help='Also generate PDF output')
@click.option('--toc/--no-toc', default=True, help='Add table of contents')
@click.option('--number/--no-number', default=True, help='Number sections')
def main(input_file, output_file, template, pdf, toc, number):
    """
    Convert Markdown to professionally formatted academic DOCX/PDF.
    
    Examples:
        python academic_md2pdf.py input.md output.docx
        python academic_md2pdf.py input.md output.docx --pdf
        python academic_md2pdf.py input.md -p --no-toc
    """
    # Determine output file
    if not output_file:
        input_path = Path(input_file)
        output_file = input_path.with_suffix('.docx')
    
    output_path = Path(output_file)
    
    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Create converter
    converter = AcademicConverter(template_file=template)
    
    print("Converting {} to {}...".format(input_file, output_file))
    print("  - Table of Contents: {}".format('Yes' if toc else 'No'))
    print("  - Section Numbering: {}".format('Yes' if number else 'No'))
    
    # Convert to DOCX
    try:
        docx_path = converter.convert_markdown_to_docx(
            input_file, 
            str(output_path),
            add_toc=toc,
            number_sections=number
        )
        
        # Convert to PDF if requested
        if pdf:
            pdf_path = output_path.with_suffix('.pdf')
            converter.auto_convert_to_pdf(docx_path, str(pdf_path))
        
        print("\n[SUCCESS] Conversion complete!")
        
    except Exception as e:
        print("\n[ERROR] Error during conversion: {}".format(e))
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()

