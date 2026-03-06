#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Thesis Markdown to Word Converter
Pre-processes Markdown tables and uses pandoc for reliable conversion.
"""

import os
import subprocess
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class ThesisMarkdownConverter:
    def __init__(self, md_file, output_dir=None):
        self.md_file = md_file
        self.output_dir = output_dir or os.path.dirname(md_file)
        
        if not self._check_pandoc():
            raise RuntimeError("pandoc is not installed. Please install it first: brew install pandoc")
    
    def _check_pandoc(self):
        """Check if pandoc is available."""
        try:
            subprocess.run(['pandoc', '--version'], capture_output=True, check=True)
            return True
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False
    
    def _preprocess_markdown(self, content):
        """Pre-process Markdown to normalize tables before pandoc conversion."""
        lines = content.split('\n')
        processed_lines = []
        i = 0
        table_num = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Detect table rows (lines starting with |)
            if line.strip().startswith('|'):
                table_num += 1
                # Extract entire table
                table_lines = [line]
                i += 1
                
                # Collect all consecutive table rows
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                
                # Count cells before normalization
                first_line = table_lines[0].strip()
                cell_count_before = first_line.count('|') - 1
                
                # Normalize table
                normalized_table = self._normalize_table(table_lines)
                
                # Count cells after normalization
                if normalized_table:
                    first_normalized = normalized_table[0]
                    cell_count_after = first_normalized.count('|') - 1
                    
                    if cell_count_before != cell_count_after:
                        print("  [Table %d] Normalized: %d cells -> %d cells" % (table_num, cell_count_before, cell_count_after))
                
                processed_lines.extend(normalized_table)
            else:
                processed_lines.append(line)
                i += 1
        
        return '\n'.join(processed_lines)
    
    def _normalize_table(self, table_lines):
        """Normalize table format to ensure proper Markdown table syntax."""
        if len(table_lines) < 2:
            return table_lines
        
        normalized = []
        
        for idx, line in enumerate(table_lines):
            # Clean up the line - remove multiple spaces
            line = line.strip()
            
            # Skip empty lines
            if not line or line == '|':
                continue
            
            # Handle multiple pipes (||, |||, etc.) by converting to single pipes
            # Replace multiple consecutive pipes with single pipe + space
            line = line.replace('|||', '|')
            line = line.replace('||', '|')
            
            # Remove leading/trailing pipes and normalize spacing
            line = line.strip('|').strip()
            
            # Split by pipe and clean each cell
            cells = [cell.strip() for cell in line.split('|')]
            
            # Remove empty cells
            cells = [cell for cell in cells if cell]
            
            # Skip if no cells found
            if not cells:
                continue
            
            # Rebuild the line with proper pipe formatting
            normalized_line = '| ' + ' | '.join(cells) + ' |'
            normalized.append(normalized_line)
            
            # After first data row, insert separator row if not present
            if idx == 0 and len(table_lines) > 1:
                next_line = table_lines[idx + 1].strip() if idx + 1 < len(table_lines) else ''
                if not self._is_separator_row(next_line):
                    separator_cells = ['---' for _ in cells]
                    separator = '| ' + ' | '.join(separator_cells) + ' |'
                    normalized.append(separator)
        
        return normalized
    
    def _is_separator_row(self, line):
        """Check if a line is a table separator row."""
        line = line.strip()
        if not line.startswith('|'):
            return False
        
        cells = line.split('|')[1:-1]
        
        for cell in cells:
            cell = cell.strip()
            # Check if cell contains only dashes, colons, and spaces
            if not re.match(r'^[\s\-:]+$', cell):
                return False
        
        return True
    
    def convert(self):
        """Perform the conversion using pandoc with pre/post-processing."""
        print("Reading markdown file...")
        
        if not os.path.exists(self.md_file):
            print("Error: File not found: " + self.md_file)
            return None
        
        # Read original content
        with open(self.md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Pre-process Markdown tables
        print("Pre-processing tables...")
        processed_content = self._preprocess_markdown(content)
        
        # Create temporary preprocessed file
        temp_md_file = os.path.join(self.output_dir, '.temp_thesis_convert.md')
        with open(temp_md_file, 'w', encoding='utf-8') as f:
            f.write(processed_content)
        
        # Generate output filename
        output_file = os.path.join(
            self.output_dir,
            Path(self.md_file).stem + '_thesis.docx'
        )
        
        print("Converting with pandoc...")
        
        # Build pandoc command
        cmd = [
            'pandoc',
            temp_md_file,
            '-o', output_file,
            '-f', 'gfm',
            '-t', 'docx',
            '--standalone',
            '--wrap=none',
            '--from', 'gfm+pipe_tables',
            '--to', 'docx',
            '-V', 'papersize=a4',
            '-V', 'geometry:margin=1in'
        ]
        
        try:
            # Execute pandoc
            result = subprocess.run(cmd, capture_output=True, text=True, check=True)
            
            # Check if file was created
            if os.path.exists(output_file):
                print("Post-processing tables and formatting...")
                self._post_process_tables(output_file)
                
                # Clean up temporary file
                if os.path.exists(temp_md_file):
                    os.remove(temp_md_file)
                
                file_size = os.path.getsize(output_file) / 1024
                print("Conversion complete!")
                print("Output: " + output_file)
                print("Size: %.1f KB" % file_size)
                return output_file
            else:
                print("Error: Output file was not created")
                return None
                
        except subprocess.CalledProcessError as e:
            print("Pandoc error: " + e.stderr)
            if os.path.exists(temp_md_file):
                os.remove(temp_md_file)
            return None
        except Exception as e:
            print("Error: " + str(e))
            if os.path.exists(temp_md_file):
                os.remove(temp_md_file)
            return None
    
    def _post_process_tables(self, docx_file):
        """Post-process tables in the DOCX file."""
        try:
            doc = Document(docx_file)
            
            table_count = len(doc.tables)
            print("Found %d table(s) to format..." % table_count)
            
            # First pass: format existing tables
            for idx, table in enumerate(doc.tables):
                self._format_table(table)
                print("Formatted table %d" % (idx + 1))
            
            # Second pass: convert text-based table blocks to proper tables
            self._convert_text_tables_to_docx(doc)
            
            # Save the modified document
            doc.save(docx_file)
            print("Tables formatted successfully")
            
        except Exception as e:
            print("Warning: Could not post-process tables: " + str(e))
    
    def _convert_text_tables_to_docx(self, doc):
        """Convert text-based table blocks to proper DOCX tables."""
        try:
            # Scan for paragraph blocks that look like tables
            i = 0
            while i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                text = para.text.strip()
                
                # Check if this paragraph looks like a table row (starts with |)
                if text.startswith('|') and '|' in text[1:]:
                    # Found potential table block, collect consecutive rows
                    table_rows = []
                    j = i
                    
                    while j < len(doc.paragraphs):
                        p = doc.paragraphs[j]
                        pt = p.text.strip()
                        
                        if pt.startswith('|') and '|' in pt[1:]:
                            table_rows.append(pt)
                            j += 1
                        elif pt == '' and j > i and j < len(doc.paragraphs) - 1:
                            # Allow one empty line
                            j += 1
                            continue
                        else:
                            break
                    
                    if len(table_rows) >= 2:
                        # Convert text rows to table format
                        self._convert_text_block_to_table(doc, i, table_rows)
                        i = j
                    else:
                        i += 1
                else:
                    i += 1
        except Exception as e:
            print("Warning: Could not convert text tables: " + str(e))
    
    def _convert_text_block_to_table(self, doc, start_idx, text_rows):
        """Convert a block of text table rows to a proper DOCX table."""
        try:
            # Parse rows into cells
            parsed_rows = []
            max_cols = 0
            
            for row_text in text_rows:
                # Parse cells from pipe-delimited text
                cells = [c.strip() for c in row_text.split('|')[1:-1]]
                cells = [c for c in cells if c]  # Remove empty cells
                
                if cells:
                    parsed_rows.append(cells)
                    max_cols = max(max_cols, len(cells))
            
            if len(parsed_rows) >= 2 and max_cols >= 2:
                # Skip separator rows
                filtered_rows = []
                for row in parsed_rows:
                    # Check if it's a separator row (all dashes)
                    if not all(re.match(r'^[\s\-:]+$', cell) for cell in row):
                        filtered_rows.append(row)
                
                if len(filtered_rows) >= 1:
                    # Find insertion point in document
                    parent = doc.paragraphs[start_idx]._element.getparent()
                    insert_idx = list(parent).index(doc.paragraphs[start_idx]._element)
                    
                    # Create new table
                    new_table = doc.add_table(rows=len(filtered_rows), cols=max_cols)
                    new_table.style = 'Light Grid Accent 1'
                    
                    # Fill table with content
                    for row_idx, row_data in enumerate(filtered_rows):
                        for col_idx, cell_content in enumerate(row_data):
                            if col_idx < max_cols:
                                cell = new_table.rows[row_idx].cells[col_idx]
                                cell.text = cell_content
                    
                    # Format the new table
                    self._format_table(new_table)
                    
                    print("Converted text block to table (rows=%d, cols=%d)" % (len(filtered_rows), max_cols))
        except Exception as e:
            print("Warning: Could not convert text block: " + str(e))
    
    def _format_table(self, table):
        """Format a single table with professional styling."""
        self._set_table_borders(table)
        self._auto_fit_columns(table)
        
        for i, row in enumerate(table.rows):
            # Set row height
            row.height = Inches(0.3)
            
            for j, cell in enumerate(row.cells):
                # Ensure cell margins
                tcPr = cell._element.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                
                for margin_name in ['top', 'left', 'bottom', 'right']:
                    margin = OxmlElement('w:%s' % margin_name)
                    margin.set(qn('w:w'), '50')
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                
                # Remove existing margins if any
                existing_mar = tcPr.find(qn('w:tcMar'))
                if existing_mar is not None:
                    tcPr.remove(existing_mar)
                tcPr.append(tcMar)
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        
                        if i == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    if i == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.left_indent = Inches(0.05)
                    paragraph.paragraph_format.right_indent = Inches(0.05)
                    paragraph.paragraph_format.line_spacing = 1.0
                
                self._set_cell_background(cell, 'FFFFFF')
    
    def _set_table_borders(self, table):
        """Apply professional solid black borders to table (no dashes)."""
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        
        tblBorders = OxmlElement('w:tblBorders')
        
        # Use solid lines for all borders (no dashes, no special effects)
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement('w:' + border_name)
            border.set(qn('w:val'), 'single')      # Single solid line (no dashes)
            border.set(qn('w:sz'), '12')           # Border size (1.5pt)
            border.set(qn('w:space'), '0')         # No space between border and cell
            border.set(qn('w:color'), '000000')    # Black color
            border.set(qn('w:shadow'), '0')        # No shadow effect
            tblBorders.append(border)
        
        tblPr.append(tblBorders)
    
    def _auto_fit_columns(self, table):
        """Auto-fit column widths to content."""
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            existing_layout = tblPr.find(qn('w:tblLayout'))
            if existing_layout is not None:
                tblPr.remove(existing_layout)
            
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'auto')
            tblPr.append(tblLayout)
        except Exception as e:
            print("Warning: Could not auto-fit columns: " + str(e))
    
    def _set_cell_background(self, cell, fill_color):
        """Set cell background color (white/no background)."""
        try:
            tcPr = cell._element.get_or_add_tcPr()
            
            existing_shade = tcPr.find(qn('w:shd'))
            if existing_shade is not None:
                tcPr.remove(existing_shade)
            
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill_color)
            tcPr.append(shd)
        except Exception as e:
            print("Warning: Could not set cell background: " + str(e))


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python thesis_md2docx.py <markdown_file> [output_dir]")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        converter = ThesisMarkdownConverter(md_file, output_dir)
        output = converter.convert()
        sys.exit(0 if output else 1)
    except RuntimeError as e:
        print("Error: " + str(e))
        sys.exit(1)
