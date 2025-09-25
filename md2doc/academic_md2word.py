#!/usr/bin/env python3
"""
Academic Markdown to Word Converter
----------------------------------
Converts Markdown files to Microsoft Word (.docx) documents with academic formatting.
Supports inline numbered points that restart from 1 in each section.
"""

import os
import sys
import re
import click
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.oxml import OxmlElement

def convert_md_to_docx(md_file, docx_file, template_file=None, img_dir=None, compact_mode=False):
    """Convert a Markdown file to a Word document with academic formatting."""
    try:
        # Create document
        if template_file and os.path.exists(template_file):
            doc = Document(template_file)
        else:
            doc = Document()
        
        # Set up document formatting
        setup_document_formatting(doc, compact_mode)
        
        # Read Markdown content
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
            
        # Extract and process references separately
        references_section = ""
        main_content = md_content
        
        # Find the References section
        ref_match = re.search(r'#+\s*References\s*\n([\s\S]+?)(?:\n#+\s+|\Z)', md_content)
        if ref_match:
            references_section = ref_match.group(1)
            # Keep the main content but remove references section
            main_content = md_content.replace(ref_match.group(0), '')
        
        # Process appendices separately (if any)
        appendix_sections = []
        appendix_regex = re.compile(r'(#+\s*Appendix\s+[A-Z].+?\n(?:[\s\S]+?)(?=\n#+\s+Appendix\s+[A-Z]|$))', re.MULTILINE)
        for appendix_match in appendix_regex.finditer(md_content):
            appendix_sections.append(appendix_match.group(1))
            # Remove the appendix section from main content
            main_content = main_content.replace(appendix_match.group(1), '')
        
        # Sort appendices by letter
        appendix_sections.sort(key=lambda x: re.search(r'#+\s*Appendix\s+([A-Z])', x).group(1))
        
        # DIRECT APPROACH: Parse markdown and build document directly
        # Process the document line by line
        lines = main_content.split('\n')
        
        # Keep track of the current section for numbering
        current_section = None
        section_counters = {}
        
        # Track when we've processed the first title to avoid duplication
        title_added = False
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Process headings directly
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2).strip()
                
                # Check if this is the title (H1) and if we've already added it
                if level == 1 and not title_added:
                    # Use Word's built-in heading style for title
                    heading = doc.add_heading(text, level=0)  # Title style
                    title_added = True
                elif level == 1 and title_added:
                    # Skip duplicate title
                    pass
                else:
                    # Use Word's built-in heading styles
                    heading = doc.add_heading(text, level=level)
                
                # Track section for numbering
                if level <= 3:
                    current_section = text
                    section_counters[current_section] = 1
            
            # Process tables
            elif line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[-:]+\|', lines[i+1]):
                # This is a markdown table
                table_rows = [line]
                i += 1  # Skip the header separator line
                table_rows.append(lines[i])
                
                # Collect all table rows
                i += 1
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_rows.append(lines[i].strip())
                    i += 1
                i -= 1  # Step back to allow the loop increment to work
                
                # Process and add the table
                create_table_from_markdown(doc, table_rows)
            
            # Process paragraphs
            elif line and not line.startswith('```'):
                # Collect all lines in the paragraph
                paragraph_text = line
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#') and not lines[i].strip().startswith('```') and not lines[i].strip().startswith('|'):
                    paragraph_text += ' ' + lines[i].strip()
                    i += 1
                i -= 1  # Step back to allow the loop increment to work
                
                # Check for custom numbered points in the paragraph
                numbered_point_match = re.match(r'^\s*(\d+)\.\s+([A-Z][a-zA-Z]+):', paragraph_text)
                if numbered_point_match and current_section in section_counters:
                    point_label = numbered_point_match.group(2)
                    point_number = section_counters[current_section]
                    
                    # Add the numbered point with the current section counter
                    paragraph = doc.add_paragraph(style='List Number')
                    run = paragraph.add_run("{0}. {1}:".format(point_number, point_label))
                    run.bold = True
                    
                    # Add the rest of the text
                    rest_of_text = paragraph_text[numbered_point_match.end():].strip()
                    if rest_of_text:
                        paragraph.add_run(' ' + rest_of_text)
                    
                    # Increment the counter for the current section
                    section_counters[current_section] += 1
                # Check for keywords section
                elif paragraph_text.lower().startswith('keywords') or paragraph_text.lower().startswith('**keywords'):
                    # Format keywords differently
                    paragraph = doc.add_paragraph(style='Normal')
                    # Extract keywords part
                    if ":" in paragraph_text:
                        keyword_prefix, keyword_list = paragraph_text.split(":", 1)
                        # Add the "Keywords:" part in bold
                        run = paragraph.add_run(keyword_prefix + ":")
                        run.bold = True
                        # Add the actual keywords
                        paragraph.add_run(" " + keyword_list.strip())
                    else:
                        paragraph.add_run(paragraph_text)
                else:
                    # Regular paragraph
                    paragraph = doc.add_paragraph(style='Normal')
                    # Format bold and italic text
                    format_markdown_text(paragraph_text, paragraph)
            
            # Process code blocks with syntax highlighting
            elif line.startswith('```'):
                # Start of a code block
                code_language = line[3:].strip().lower()  # Extract language if specified
                code_block_content = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_block_content.append(lines[i])
                    i += 1
                
                # Add code block with potential syntax highlighting
                for code_line in code_block_content:
                    paragraph = doc.add_paragraph(style='Code Block')
                    
                    # Simple syntax highlighting for Python
                    if code_language == 'python':
                        # Apply basic Python syntax highlighting
                        parts = []
                        # Keywords
                        keywords = ['def', 'class', 'import', 'from', 'return', 'if', 'else', 'elif',
                                   'for', 'while', 'try', 'except', 'finally', 'with', 'as', 'in', 'not',
                                   'and', 'or', 'True', 'False', 'None', 'async', 'await', 'yield']
                        
                        # Split the line to identify different components
                        processed = False
                        
                        # Check for comments
                        if '#' in code_line:
                            comment_idx = code_line.find('#')
                            code_part = code_line[:comment_idx]
                            comment_part = code_line[comment_idx:]
                            
                            # Add the code part with normal processing
                            if code_part:
                                run = paragraph.add_run(code_part)
                            
                            # Add the comment with green color
                            comment_run = paragraph.add_run(comment_part)
                            comment_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for comments
                            processed = True
                        
                        # Check for string literals
                        elif not processed and ('"' in code_line or "'" in code_line):
                            # Very basic string detection - not handling all cases
                            parts = re.split(r'(".*?"|\'.*?\')', code_line)
                            for idx, part in enumerate(parts):
                                if idx % 2 == 1:  # This is a string literal
                                    run = paragraph.add_run(part)
                                    run.font.color.rgb = RGBColor(163, 21, 21)  # Brown/Red for strings
                                else:
                                    run_processed = False
                                    for word in re.findall(r'\b\w+\b', part):
                                        if word in keywords:
                                            idx = part.find(word)
                                            prefix = part[:idx]
                                            suffix = part[idx + len(word):]
                                            
                                            if prefix:
                                                paragraph.add_run(prefix)
                                            
                                            keyword_run = paragraph.add_run(word)
                                            keyword_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for keywords
                                            
                                            if suffix:
                                                paragraph.add_run(suffix)
                                            
                                            run_processed = True
                                    
                                    if not run_processed:
                                        paragraph.add_run(part)
                            processed = True
                        
                        # Just check for keywords
                        if not processed:
                            current_idx = 0
                            for keyword in keywords:
                                pattern = r'\b' + re.escape(keyword) + r'\b'
                                for match in re.finditer(pattern, code_line):
                                    start, end = match.span()
                                    
                                    # Add text before the keyword
                                    if start > current_idx:
                                        paragraph.add_run(code_line[current_idx:start])
                                    
                                    # Add the keyword with blue color
                                    keyword_run = paragraph.add_run(keyword)
                                    keyword_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for keywords
                                    
                                    current_idx = end
                            
                            # Add any remaining text
                            if current_idx < len(code_line):
                                paragraph.add_run(code_line[current_idx:])
                    else:
                        # For other languages, just apply monospace formatting
                        paragraph.add_run(code_line)
            
            i += 1
        
        # Add References section if found
        if references_section:
            # Add References heading
            doc.add_heading("References", level=2)
            
            # Process each reference line
            ref_lines = references_section.strip().split('\n')
            for ref_line in ref_lines:
                if ref_line.strip():
                    # Format reference with proper style and hanging indent
                    paragraph = doc.add_paragraph(style='Reference Style')
                    format_markdown_text(ref_line.strip(), paragraph)
                    
                    # Ensure hanging indent
                    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
                    paragraph.paragraph_format.left_indent = Inches(0.5)
        
        # Process each appendix section
        for appendix_content in appendix_sections:
            process_appendix_section(appendix_content, doc)
        
        # Save document
        doc.save(docx_file)
        print("Successfully converted {0} to {1} with academic formatting".format(md_file, docx_file))
        return True
    
    except Exception as e:
        print("Error converting {0} to {1}: {2}".format(md_file, docx_file, e))
        import traceback
        traceback.print_exc()
        return False

def create_table_from_markdown(doc, table_rows):
    """Convert markdown table to Word table."""
    if not table_rows or len(table_rows) < 2:
        return
    
    # Process header row to determine column count
    header_row = table_rows[0].strip('|').split('|')
    col_count = len(header_row)
    row_count = len(table_rows) - 1  # Subtract 1 for the separator row
    
    # Get alignment info from separator row
    alignments = []
    sep_cells = table_rows[1].strip('|').split('|')
    for cell in sep_cells:
        cell = cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif cell.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
    
    # Create the table
    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = 'Table Grid'
    
    # Add header row
    for i, cell_text in enumerate(header_row):
        if i < len(table.rows[0].cells):
            cell = table.rows[0].cells[i]
            cell.text = cell_text.strip()
            
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                if i < len(alignments):
                    paragraph.alignment = alignments[i]
    
    # Add data rows
    for row_idx in range(2, len(table_rows)):
        doc_row_idx = row_idx - 2
        if doc_row_idx >= len(table.rows):
            continue
            
        cells = table_rows[row_idx].strip('|').split('|')
        for col_idx, cell_text in enumerate(cells):
            if col_idx < len(table.rows[doc_row_idx].cells):
                cell = table.rows[doc_row_idx].cells[col_idx]
                cell.text = cell_text.strip()
                
                # Apply alignment
                if col_idx < len(alignments):
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = alignments[col_idx]
    
    # Add spacing after table
    doc.add_paragraph()

def process_appendix_section(appendix_content, doc):
    """Process an appendix section and add it to the document."""
    lines = appendix_content.split('\n')
    
    in_code_block = False
    code_language = None
    code_block_content = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if in_code_block:
            if line.startswith('```'):
                # End of code block
                in_code_block = False
                
                # Process the collected code block
                for code_line in code_block_content:
                    paragraph = doc.add_paragraph(style='Code Block')
                    
                    # Simple syntax highlighting for Python
                    if code_language == 'python':
                        # Apply basic Python syntax highlighting
                        parts = []
                        # Keywords
                        keywords = ['def', 'class', 'import', 'from', 'return', 'if', 'else', 'elif',
                                   'for', 'while', 'try', 'except', 'finally', 'with', 'as', 'in', 'not',
                                   'and', 'or', 'True', 'False', 'None', 'async', 'await', 'yield']
                        
                        # Split the line to identify different components
                        processed = False
                        
                        # Check for comments
                        if '#' in code_line:
                            comment_idx = code_line.find('#')
                            code_part = code_line[:comment_idx]
                            comment_part = code_line[comment_idx:]
                            
                            # Add the code part with normal processing
                            if code_part:
                                run = paragraph.add_run(code_part)
                            
                            # Add the comment with green color
                            comment_run = paragraph.add_run(comment_part)
                            comment_run.font.color.rgb = RGBColor(0, 128, 0)  # Green for comments
                            processed = True
                        
                        # Check for string literals
                        elif not processed and ('"' in code_line or "'" in code_line):
                            # Very basic string detection - not handling all cases
                            parts = re.split(r'(".*?"|\'.*?\')', code_line)
                            for idx, part in enumerate(parts):
                                if idx % 2 == 1:  # This is a string literal
                                    run = paragraph.add_run(part)
                                    run.font.color.rgb = RGBColor(163, 21, 21)  # Brown/Red for strings
                                else:
                                    run_processed = False
                                    for word in re.findall(r'\b\w+\b', part):
                                        if word in keywords:
                                            idx = part.find(word)
                                            prefix = part[:idx]
                                            suffix = part[idx + len(word):]
                                            
                                            if prefix:
                                                paragraph.add_run(prefix)
                                            
                                            keyword_run = paragraph.add_run(word)
                                            keyword_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for keywords
                                            
                                            if suffix:
                                                paragraph.add_run(suffix)
                                            
                                            run_processed = True
                                    
                                    if not run_processed:
                                        paragraph.add_run(part)
                            processed = True
                        
                        # Just check for keywords
                        if not processed:
                            current_idx = 0
                            for keyword in keywords:
                                pattern = r'\b' + re.escape(keyword) + r'\b'
                                for match in re.finditer(pattern, code_line):
                                    start, end = match.span()
                                    
                                    # Add text before the keyword
                                    if start > current_idx:
                                        paragraph.add_run(code_line[current_idx:start])
                                    
                                    # Add the keyword with blue color
                                    keyword_run = paragraph.add_run(keyword)
                                    keyword_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for keywords
                                    
                                    current_idx = end
                            
                            # Add any remaining text
                            if current_idx < len(code_line):
                                paragraph.add_run(code_line[current_idx:])
                    else:
                        # For other languages, just apply monospace formatting
                        paragraph.add_run(code_line)
                
                code_block_content = []
            else:
                # Add line to code block
                code_block_content.append(line)
        else:
            # Check if this line starts a code block
            if line.startswith('```'):
                in_code_block = True
                code_language = line[3:].strip().lower()
            # Process tables in appendix
            elif line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[-:]+\|', lines[i+1]):
                # This is a markdown table
                table_rows = [line]
                i += 1  # Skip the header separator line
                table_rows.append(lines[i])
                
                # Collect all table rows
                i += 1
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_rows.append(lines[i].strip())
                    i += 1
                i -= 1  # Step back to allow the loop increment to work
                
                # Process and add the table
                create_table_from_markdown(doc, table_rows)
            # Process headings
            else:
                header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if header_match:
                    level = len(header_match.group(1))
                    text = header_match.group(2).strip()
                    # Create heading using Word's built-in heading styles
                    # For appendix, we still want to use built-in heading styles
                    heading = doc.add_heading(text, level=level)
                
                # Process paragraphs
                elif line and not line.startswith('```'):
                    # Collect all lines in the paragraph
                    paragraph_text = line
                    j = i + 1
                    while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith('#') and not lines[j].strip().startswith('```') and not lines[j].strip().startswith('|'):
                        paragraph_text += ' ' + lines[j].strip()
                        j += 1
                    i = j - 1  # Update the counter to skip processed lines
                    
                    # Regular paragraph
                    paragraph = doc.add_paragraph(style='Normal')
                    # Format bold and italic text
                    format_markdown_text(paragraph_text, paragraph)
        
        i += 1

def setup_document_formatting(doc, compact_mode=False):
    """Set up academic document formatting with optional space-saving settings."""
    # Create custom styles for different document elements
    create_custom_styles(doc, compact_mode)
    
    # Set document margins
    if compact_mode:
        # Set document margins (2 cm = ~0.79 inches)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.79)    # 2 cm top margin
            section.bottom_margin = Inches(0.79) # 2 cm bottom margin
            section.left_margin = Inches(0.79)   # 2 cm left margin
            section.right_margin = Inches(0.79)  # 2 cm right margin
            
            # Header and footer settings
            section.header_distance = Inches(0.4)  # 1.0 cm header height
            section.footer_distance = Inches(0.4)  # 1.0 cm footer height

def create_custom_styles(doc, compact_mode=False):
    """Create custom styles for different document elements."""
    
    # Base font settings based on mode
    base_font = 'Arial' if compact_mode else 'Times New Roman'
    base_font_size = Pt(11) if compact_mode else Pt(12)
    line_spacing = 1.15 if compact_mode else 1.5
    
    # No need to create custom heading styles since we're using Word's built-in ones
    # Just customize other styles we need
    
    # Code block style
    try:
        code_style = doc.styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9) if compact_mode else Pt(10)
        code_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        code_style.paragraph_format.line_spacing = 1.0  # Single spacing for code
        code_style.paragraph_format.space_before = Pt(3) if compact_mode else Pt(6)
        code_style.paragraph_format.space_after = Pt(3) if compact_mode else Pt(6)
        code_style.paragraph_format.left_indent = Inches(0.2) if compact_mode else Inches(0.5)
        code_style.paragraph_format.right_indent = Inches(0.2) if compact_mode else Inches(0.5)
        code_style.paragraph_format.first_line_indent = Inches(0)
        # Add background shading for code blocks
        code_style.paragraph_format.keep_together = True
        if hasattr(code_style, 'paragraph_format') and hasattr(code_style.paragraph_format, 'shading'):
            code_style.paragraph_format.shading.background_pattern_color = RGBColor(245, 245, 245)
    except ValueError:
        # Style already exists
        pass
    
    # Blockquote style
    try:
        quote_style = doc.styles.add_style('Block Quote', WD_STYLE_TYPE.PARAGRAPH)
        quote_style.base_style = doc.styles['Normal']
        quote_style.font.italic = True
        quote_style.paragraph_format.left_indent = Inches(0.4) if compact_mode else Inches(1.0)
        quote_style.paragraph_format.right_indent = Inches(0.2) if compact_mode else Inches(0.5)
        quote_style.paragraph_format.space_before = Pt(3) if compact_mode else Pt(6)
        quote_style.paragraph_format.space_after = Pt(3) if compact_mode else Pt(6)
        quote_style.paragraph_format.first_line_indent = Inches(0)
    except ValueError:
        # Style already exists
        pass
    
    # Reference style - enhance with better formatting
    try:
        reference_style = doc.styles.add_style('Reference Style', WD_STYLE_TYPE.PARAGRAPH)
        reference_style.base_style = doc.styles['Normal']
        reference_style.font.size = Pt(10) if compact_mode else Pt(11)
        reference_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        reference_style.paragraph_format.left_indent = Inches(0.5)
        reference_style.paragraph_format.first_line_indent = Inches(-0.5)  # Hanging indent
        reference_style.paragraph_format.space_before = Pt(0)
        reference_style.paragraph_format.space_after = Pt(6)
        reference_style.paragraph_format.line_spacing = 1.0  # Single spacing for references
    except ValueError:
        # Style already exists
        pass

def format_markdown_text(text, paragraph):
    """Format markdown text with bold, italic formatting."""
    # Replace markdown formatting with Word formatting
    
    # Split text by potential formatting markers
    segments = re.split(r'(\*\*|\*|__|\b_\b)', text)
    
    # Track formatting state
    bold_active = False
    italic_active = False
    
    # Process each segment
    current_text = ""
    
    for segment in segments:
        if segment == "**" or segment == "__":
            # Toggle bold
            if current_text:
                run = paragraph.add_run(current_text)
                run.bold = bold_active
                run.italic = italic_active
                current_text = ""
            bold_active = not bold_active
        elif segment == "*" or segment == "_":
            # Toggle italic
            if current_text:
                run = paragraph.add_run(current_text)
                run.bold = bold_active
                run.italic = italic_active
                current_text = ""
            italic_active = not italic_active
        else:
            # Regular text
            current_text += segment
    
    # Add any remaining text
    if current_text:
        run = paragraph.add_run(current_text)
        run.bold = bold_active
        run.italic = italic_active

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Top
    if top:
        tcBorders = tcPr.first_child_found_or_add(qn('w:tcBorders'))
        topElem = OxmlElement('w:top')
        topElem.set(qn('w:val'), 'single')
        topElem.set(qn('w:sz'), str(top))
        topElem.set(qn('w:space'), '0')
        topElem.set(qn('w:color'), 'auto')
        tcBorders.append(topElem)
    
    # Bottom
    if bottom:
        tcBorders = tcPr.first_child_found_or_add(qn('w:tcBorders')) 
        bottomElem = OxmlElement('w:bottom')
        bottomElem.set(qn('w:val'), 'single')
        bottomElem.set(qn('w:sz'), str(bottom))
        bottomElem.set(qn('w:space'), '0')
        bottomElem.set(qn('w:color'), 'auto')
        tcBorders.append(bottomElem)
        
    # Left
    if left:
        tcBorders = tcPr.first_child_found_or_add(qn('w:tcBorders'))
        leftElem = OxmlElement('w:left')
        leftElem.set(qn('w:val'), 'single')
        leftElem.set(qn('w:sz'), str(left))
        leftElem.set(qn('w:space'), '0')
        leftElem.set(qn('w:color'), 'auto')
        tcBorders.append(leftElem)
        
    # Right
    if right:
        tcBorders = tcPr.first_child_found_or_add(qn('w:tcBorders'))
        rightElem = OxmlElement('w:right')
        rightElem.set(qn('w:val'), 'single')
        rightElem.set(qn('w:sz'), str(right))
        rightElem.set(qn('w:space'), '0')
        rightElem.set(qn('w:color'), 'auto')
        tcBorders.append(rightElem)

@click.command()
@click.argument('input_file', type=click.Path(exists=True))
@click.argument('output_file', type=click.Path(), required=False)
@click.option('--template', '-t', type=click.Path(exists=True), help='Word template to use as base document')
@click.option('--img-dir', '-i', type=click.Path(exists=True), help='Directory containing images referenced in Markdown')
@click.option('--output-dir', '-o', type=click.Path(exists=True), help='Directory to save output files (default: docoutput)', default='docoutput')
@click.option('--edit-original', '-e', is_flag=True, help='Edit the original Markdown file to update numbering')
@click.option('--compact', '-c', is_flag=True, help='Use space-saving compact formatting')
def main(input_file, output_file, template, img_dir, output_dir, edit_original, compact):
    """Convert a Markdown file to a Word document with academic formatting.
    
    INPUT_FILE is the path to the Markdown file to convert.
    
    OUTPUT_FILE is the optional path to save the Word document. If not provided,
    the output will be saved in the output directory with the same name as the input file but with a .docx extension.
    """
    if edit_original:
        # Read the original file
        with open(input_file, 'r', encoding='utf-8') as f:
            original_content = f.read()
        
        # Edit the file to fix inline numbering
        edited_content = edit_markdown_numbering(original_content)
        
        # Save the edited content back to the file
        with open(input_file, 'w', encoding='utf-8') as f:
            f.write(edited_content)
        
        print("Updated inline numbering in {0}".format(input_file))
    
    if not output_file:
        # Generate output file name based on input file
        input_basename = os.path.basename(input_file)
        input_name, _ = os.path.splitext(input_basename)
        output_file = os.path.join(output_dir, "{0}.docx".format(input_name))
    
    # Create output directory if it doesn't exist
    output_dir_path = os.path.dirname(output_file)
    if output_dir_path and not os.path.exists(output_dir_path):
        os.makedirs(output_dir_path)
    
    # Convert the file
    success = convert_md_to_docx(input_file, output_file, template, img_dir, compact)
    
    if success:
        sys.exit(0)
    else:
        sys.exit(1)

def edit_markdown_numbering(content):
    """
    Edit the original markdown file to update inline numbering.
    This will restart numbered points from 1 in each section.
    """
    # Split content by headings (## or ### level headings)
    sections = re.split(r'^(#{2,3}\s+.+)$', content, flags=re.MULTILINE)
    
    if len(sections) <= 1:
        return content
    
    result = []
    for i in range(0, len(sections), 2):
        if i == 0:
            # This is content before the first heading
            section_content = sections[i]
            result.append(section_content)
        
        if i + 1 < len(sections):
            # This is a heading followed by content
            heading = sections[i]
            section_content = sections[i+1] if i+1 < len(sections) else ""
            
            # Add the heading
            result.append(heading)
            
            # Find numbered points like "16. Retriever:" and update them
            counter = 1
            def replace_numbering(match):
                nonlocal counter
                num = counter
                counter += 1
                return "{0}. {1}:".format(num, match.group(2))
            
            updated_content = re.sub(
                r'(\d+)\.\s+([A-Z][a-zA-Z]+):',
                replace_numbering, 
                section_content
            )
            
            result.append(updated_content)
    
    return ''.join(result)

if __name__ == '__main__':
    main()
